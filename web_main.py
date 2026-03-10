import os
import json
from flask import Flask, request, jsonify, render_template, redirect
from flask_cors import CORS
from dotenv import load_dotenv
import requests as http_requests

load_dotenv()

# Templates at root level - most reliable on Render
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
app = Flask(__name__, 
            template_folder=BASE_DIR,
            static_folder=os.path.join(BASE_DIR, 'static'))
CORS(app)

GROQ_API_KEY = os.environ.get("GROQ_API_KEY", "")
SUPABASE_URL = os.environ.get("SUPABASE_URL", "")
SUPABASE_KEY = os.environ.get("SUPABASE_KEY", "")

print(f"[BOOT] GROQ_API_KEY={'SET' if GROQ_API_KEY else 'MISSING'}")
print(f"[BOOT] SUPABASE_URL={'SET' if SUPABASE_URL else 'MISSING'} ({SUPABASE_URL[:30]}...)" if SUPABASE_URL else "[BOOT] SUPABASE_URL=MISSING")
print(f"[BOOT] SUPABASE_KEY={'SET' if SUPABASE_KEY else 'MISSING'}")

_supabase_error = None


# ---------------------------------------------------------------------------
#  Lightweight Supabase REST wrapper (replaces supabase-py SDK)
#  Uses PostgREST endpoints directly so any API key format works.
# ---------------------------------------------------------------------------

class _SupabaseResponse:
    """Mimics supabase-py execute() result with .data and .count."""
    def __init__(self, data=None, count=None):
        self.data = data if data is not None else []
        self.count = count


class _QueryBuilder:
    """Chainable PostgREST query builder."""

    def __init__(self, base_url, table, headers):
        self._url = f"{base_url}/rest/v1/{table}"
        self._headers = dict(headers)
        self._params = {}
        self._method = "GET"
        self._body = None
        self._count_mode = None

    # --- column selection ---
    def select(self, columns="*", count=None):
        self._method = "GET"
        self._params["select"] = columns
        if count:
            self._count_mode = count          # "exact", "planned", "estimated"
        return self

    # --- filters ---
    def eq(self, column, value):
        self._params[column] = f"eq.{value}"
        return self

    def neq(self, column, value):
        self._params[column] = f"neq.{value}"
        return self

    def gt(self, column, value):
        self._params[column] = f"gt.{value}"
        return self

    def lt(self, column, value):
        self._params[column] = f"lt.{value}"
        return self

    # --- modifiers ---
    def order(self, column, desc=True):
        direction = "desc" if desc else "asc"
        self._params["order"] = f"{column}.{direction}"
        return self

    def limit(self, n):
        self._params["limit"] = str(n)
        return self

    # --- mutations ---
    def upsert(self, data, on_conflict=None):
        self._method = "POST"
        self._headers["Prefer"] = "resolution=merge-duplicates,return=representation"
        if on_conflict:
            self._params["on_conflict"] = on_conflict
        self._body = data
        return self

    def insert(self, data):
        self._method = "POST"
        self._headers["Prefer"] = "return=representation"
        self._body = data
        return self

    def update(self, data):
        self._method = "PATCH"
        self._headers["Prefer"] = "return=representation"
        self._body = data
        return self

    def delete(self):
        self._method = "DELETE"
        self._headers["Prefer"] = "return=representation"
        return self

    # --- execute ---
    def execute(self):
        headers = dict(self._headers)
        if self._count_mode:
            headers["Prefer"] = headers.get("Prefer", "")
            if headers["Prefer"]:
                headers["Prefer"] += f",count={self._count_mode}"
            else:
                headers["Prefer"] = f"count={self._count_mode}"

        if self._body is not None:
            headers["Content-Type"] = "application/json"

        resp = http_requests.request(
            method=self._method,
            url=self._url,
            headers=headers,
            params=self._params,
            json=self._body if self._body is not None else None,
            timeout=30,
        )

        # Parse count from content-range header (e.g. "0-9/42")
        count = None
        cr = resp.headers.get("content-range", "")
        if "/" in cr:
            try:
                count = int(cr.split("/")[1])
            except (ValueError, IndexError):
                pass

        # PostgREST returns [] on success for DELETE, or the rows
        try:
            data = resp.json() if resp.text else []
        except Exception:
            data = []

        # Raise on HTTP errors (4xx/5xx) so callers' except blocks catch them
        if resp.status_code >= 400:
            msg = data.get("message", resp.text[:200]) if isinstance(data, dict) else resp.text[:200]
            raise Exception(f"PostgREST {resp.status_code}: {msg}")

        return _SupabaseResponse(data=data, count=count)


class _SupabaseREST:
    """Drop-in replacement for supabase-py Client with .table() interface."""

    def __init__(self, url, key):
        self._url = url.rstrip("/")
        self._headers = {
            "apikey": key,
            "Authorization": f"Bearer {key}",
        }

    def table(self, name):
        return _QueryBuilder(self._url, name, self._headers)


def get_supabase():
    global _supabase_error
    if not SUPABASE_URL or not SUPABASE_KEY:
        _supabase_error = f"Missing env: URL={'SET' if SUPABASE_URL else 'EMPTY'}, KEY={'SET' if SUPABASE_KEY else 'EMPTY'}"
        print(f"[Supabase] {_supabase_error}")
        return None
    try:
        client = _SupabaseREST(SUPABASE_URL, SUPABASE_KEY)
        # Quick connectivity test — just validate the key works
        _supabase_error = None
        return client
    except Exception as e:
        _supabase_error = f"REST client init error: {type(e).__name__}: {e}"
        print(f"[Supabase] {_supabase_error}")
        return None


def call_claude(prompt, max_tokens=4096, model=None):
    """Call GROQ API (OpenAI-compatible). Tries llama-3.3-70b-versatile first,
    falls back to llama-3.1-8b-instant if unavailable."""
    if not GROQ_API_KEY:
        return "Error: GROQ_API_KEY not set. Add it in Render → Environment Variables."

    MODELS = [model] if model else ["llama-3.3-70b-versatile", "llama-3.1-8b-instant"]

    for m in MODELS:
        try:
            res = http_requests.post(
                "https://api.groq.com/openai/v1/chat/completions",
                headers={
                    "Authorization": f"Bearer {GROQ_API_KEY}",
                    "Content-Type": "application/json"
                },
                json={
                    "model": m,
                    "max_tokens": max_tokens,
                    "temperature": 0.3,
                    "messages": [{"role": "user", "content": prompt}]
                },
                timeout=120
            )
            print(f"[Groq] model={m} HTTP {res.status_code}, response length: {len(res.text)}")
            if res.status_code == 429:
                print(f"[Groq] Rate limit on {m}, trying next model...")
                continue
            if res.status_code != 200:
                print(f"[Groq] HTTP {res.status_code} on {m}: {res.text[:300]}")
                # Try next model on 4xx (model may be unavailable)
                if res.status_code in (400, 404, 422):
                    continue
                return f"Error: Groq HTTP {res.status_code}: {res.text[:300]}"
            data = res.json()
            if "error" in data:
                err_msg = data["error"].get("message", str(data["error"]))
                print(f"[Groq] API error on {m}: {err_msg}")
                continue  # try next model
            content = data["choices"][0]["message"]["content"]
            finish = data["choices"][0].get("finish_reason", "unknown")
            print(f"[Groq] model={m} finish_reason={finish}, content length={len(content)}")
            if finish == "length":
                print("[Groq] WARNING: response truncated — increase max_tokens or reduce prompt")
            return content
        except Exception as e:
            print(f"[Groq] Exception on {m}: {e}")
            continue

    return "Error: All Groq models failed — check API key and quota"

PROFILE = {
    "name": "Amretha Karthikeyan",
    "address": "#02-321 153 Gangsa Road, Singapore-670153",
    "mobile": "+65-90256503",
    "email": "amretha.ammu@gmail.com",
    "linkedin": "https://www.linkedin.com/in/amretha-nishanth-534b39101/",
    "headline": "Product Owner | Lead BA | Fintech & Digital Products · Singapore",
    "aiProjectUrl": "https://stock-monitor-8ak6.onrender.com",
    "summary": (
        "SAFe 6.0 certified Product Owner and Lead Business Analyst with 5+ years owning "
        "product backlogs and driving digital product delivery in fintech and banking. "
        "At KPMG Singapore, served as de-facto Product Owner for Loan IQ — a core banking "
        "platform — leading cross-functional squads (engineering, UX, QA) to ship features "
        "and deliver measurable business outcomes. Built and deployed a live AI-powered Trade "
        "Analysis platform using Claude Opus 4.6. Seeking in-house product roles to own "
        "roadmaps end-to-end, from discovery through to scale."
    ),
    "skills": [
        "Tableau", "Power BI", "PSQL", "Python", "Agile", "JIRA", "Excel",
        "Microsoft Project", "Product Vision", "Roadmapping", "Business Analysis",
        "Risk Mitigation", "Change Management", "Budget Forecasting", "Variance Analysis",
        "KPI Tracking", "Dashboard Reporting", "SAFe 6.0", "API integrations",
        "Loan IQ", "SQL", "Stakeholder Management", "Generative AI", "LLM",
        "Claude API", "AI product development", "Prompt Engineering"
    ],
    "certification": "Scaled Agile Framework 6.0 Product Owner/Product Management",
    "experience": [
        {
            "company": "KPMG, Singapore",
            "role": "Lead Business Analyst – Functional Consultant – Loan IQ",
            "period": "Feb 2021 – Present",
            "bullets": [
                "Served as de-facto Product Owner for Loan IQ core banking platform, owning the product backlog and driving sprint delivery for a cross-functional squad (engineering, UX, QA)",
                "Partnered with Enterprise Singapore on large-scale digital transformation projects",
                "Drove product scope decisions through impact analysis, generating ~5% additional business value",
                "Identified and delivered automation of interest computation workflow, eliminating 30 man-days of manual effort",
                "Owned and prioritised product backlog, ensuring alignment with business objectives and regulatory requirements",
                "Led sprint ceremonies (planning, reviews, retros, PI Planning) across multi-squad programme",
                "Managed 3rd party vendors, conducted go-live planning, and led data migrations from legacy systems",
                "Designed and executed end-to-end test scenarios on Loan IQ applications (M&A, Trade, WCL, FA)"
            ],
            "achievements": [
                "Drove ~5% business value through product scope and change request impact analysis",
                "Eliminated 30 man-days of manual work through automated interest computation feature",
                "Led team through critical sprint-to-SIT transition, maintaining delivery timeline"
            ]
        },
        {
            "company": "J.P. Morgan",
            "role": "Asset Management Virtual Internship",
            "period": "Oct 2023 – Jan 2024",
            "bullets": [
                "Gathered product requirements from trading/execution teams to build robust investor profiles",
                "Performed quantitative analysis of 5 stocks and recommended to 2 clients based on risk metrics",
                "Measured portfolio performance via KPIs: Annual Return, Portfolio Variance, Standard Deviation"
            ]
        },
        {
            "company": "Amazon Inc, India",
            "role": "Business Analyst",
            "period": "Mar 2018 – Mar 2019",
            "bullets": [
                "Built real-time quality monitoring dashboards using Power BI from SQL Server and MS Excel",
                "Translated business requirements into functional and non-functional specifications",
                "Analysed and visualised operational data using Tableau and Power BI"
            ]
        }
    ],
    "education": [
        {"degree": "Master of Science – Engineering Business Management", "school": "Coventry University, UK", "period": "Jul 2019 – Nov 2020"},
        {"degree": "Bachelor of Engineering – Electronics & Communication", "school": "Anna University, India", "period": "Jul 2012 – Jun 2016"}
    ],
    "projects": [
        {
            "title": "AI-Powered Trade Analysis Platform",
            "type": "Personal Project",
            "period": "2025",
            "url": "https://stock-monitor-8ak6.onrender.com",
            "tech": "Claude Opus 4.6 (Anthropic), Python, Flask, Render",
            "bullets": [
                "Designed and deployed a live AI-powered Trade Analysis platform using Claude Opus 4.6 — accessible at https://stock-monitor-8ak6.onrender.com",
                "Combined financial trade data and international trade flow analysis using generative AI",
                "Demonstrated end-to-end AI product development: problem definition, prompt engineering, LLM integration, Flask backend, and Render deployment",
                "Independently shipped a working AI product — demonstrating product ownership beyond theory"
            ]
        }
    ]
}

PRODUCT_FRAMING = """
CRITICAL POSITIONING — The candidate is transitioning from CONSULTING to IN-HOUSE PRODUCT roles:
- Reframe consulting experience → "Product Owner for product squad"
- Reframe "client delivery" → "shipped product features, owned backlog, drove sprint outcomes"
- DO NOT use: consultant, client, engagement, billable, service delivery
- DO USE: product, squad, roadmap, discovery, iteration, user value, outcome, feature, backlog
"""


# ─── MULTI-USER PROFILE SUPPORT ──────────────────────────────
# Amretha's profile is hardcoded as DEFAULT_PROFILE.
# Other users can upload their own profile via /api/profile/save.
# get_active_profile() returns user-uploaded profile if it exists, else default.

DEFAULT_PROFILE = PROFILE  # alias for clarity

def get_active_profile():
    """Return user-uploaded profile from Supabase, or fall back to hardcoded DEFAULT_PROFILE."""
    try:
        sb = get_supabase()
        if sb:
            res = sb.table("settings").select("value").eq("key", "user_profile").execute()
            if res.data and res.data[0].get("value"):
                custom = json.loads(res.data[0]["value"])
                if custom.get("name"):  # valid profile must have a name
                    return custom
    except Exception as e:
        print(f"[Profile] Error loading custom profile: {e}")
    return DEFAULT_PROFILE


def build_product_framing(profile):
    """Generate dynamic positioning text based on user profile."""
    name = profile.get("name", "the candidate")
    exp = profile.get("experience", [])
    current = exp[0] if exp else {}
    current_company = current.get("company", "their current company")
    return f"""
CRITICAL POSITIONING — {name} is transitioning from CONSULTING to IN-HOUSE PRODUCT roles:
- Reframe "{current_company} consultant" → "Product Owner for product squad"
- Reframe "client delivery" → "shipped product features, owned backlog, drove sprint outcomes"
- DO NOT use: consultant, client, engagement, billable, service delivery
- DO USE: product, squad, roadmap, discovery, iteration, user value, outcome, feature, backlog
"""


@app.route("/api/profile/save", methods=["POST"])
def save_profile():
    """Save a user-uploaded profile. Expects JSON with profile fields."""
    data = request.json or {}
    if not data.get("name"):
        return jsonify({"error": "Profile must include at least a name"}), 400

    # Normalize the profile structure
    profile = {
        "name":          data.get("name", "").strip(),
        "address":       data.get("address", "").strip(),
        "mobile":        data.get("mobile", "").strip(),
        "email":         data.get("email", "").strip(),
        "linkedin":      data.get("linkedin", "").strip(),
        "headline":      data.get("headline", "").strip(),
        "aiProjectUrl":  data.get("aiProjectUrl", "").strip(),
        "summary":       data.get("summary", "").strip(),
        "skills":        data.get("skills", []),
        "certification": data.get("certification", "").strip(),
        "experience":    data.get("experience", []),
        "education":     data.get("education", []),
        "projects":      data.get("projects", []),
    }

    # If skills came as a comma-separated string, split it
    if isinstance(profile["skills"], str):
        profile["skills"] = [s.strip() for s in profile["skills"].split(",") if s.strip()]

    try:
        sb = get_supabase()
        if not sb:
            return jsonify({"error": "Supabase not configured"}), 400
        sb.table("settings").upsert({"key": "user_profile", "value": json.dumps(profile)}, on_conflict="key").execute()
        return jsonify({"ok": True, "message": f"Profile saved for {profile['name']}"})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/profile/load", methods=["GET"])
def load_profile():
    """Load the active profile (custom or default)."""
    profile = get_active_profile()
    is_default = (profile.get("name") == DEFAULT_PROFILE.get("name") and
                  profile.get("email") == DEFAULT_PROFILE.get("email"))
    return jsonify({"profile": profile, "is_default": is_default})


@app.route("/api/profile/reset", methods=["POST"])
def reset_profile():
    """Reset to default profile (Amretha's hardcoded profile)."""
    try:
        sb = get_supabase()
        if sb:
            sb.table("settings").delete().eq("key", "user_profile").execute()
        return jsonify({"ok": True, "message": "Reset to default profile"})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/profile/parse-resume", methods=["POST"])
def parse_resume_to_profile():
    """Use AI to parse a pasted resume into structured profile JSON."""
    data = request.json or {}
    resume_text = data.get("resumeText", "").strip()
    if not resume_text or len(resume_text) < 50:
        return jsonify({"error": "Please paste a resume with at least 50 characters"}), 400

    prompt = f"""Parse this resume into a structured JSON profile. Extract all information accurately.

RESUME TEXT:
{resume_text[:5000]}

Return ONLY valid JSON with this exact structure (no markdown, no extra text):
{{
  "name": "Full Name",
  "address": "Address if mentioned",
  "mobile": "Phone number",
  "email": "Email address",
  "linkedin": "LinkedIn URL",
  "headline": "Professional headline (e.g., 'Product Manager | Fintech | Singapore')",
  "aiProjectUrl": "Any project/portfolio URL mentioned",
  "summary": "Professional summary (2-3 sentences)",
  "skills": ["Skill 1", "Skill 2", "Skill 3"],
  "certification": "Certifications listed, comma separated",
  "experience": [
    {{
      "company": "Company Name",
      "role": "Job Title",
      "period": "Start – End",
      "bullets": ["Achievement/responsibility 1", "Achievement 2"],
      "achievements": ["Key achievement 1"]
    }}
  ],
  "education": [
    {{"degree": "Degree Name", "school": "University Name", "period": "Start – End"}}
  ],
  "projects": [
    {{
      "title": "Project Name",
      "type": "Project type",
      "period": "Year",
      "url": "URL if any",
      "tech": "Technologies used",
      "bullets": ["Description 1"]
    }}
  ]
}}

If a field is not found in the resume, use empty string "" or empty array [].
Extract ALL experience entries, education, and skills mentioned."""

    result = call_claude(prompt)
    try:
        import re
        clean = re.sub(r'```json|```', '', result).strip()
        m = re.search(r'\{.*\}', clean, re.DOTALL)
        if m:
            profile = json.loads(m.group())
            return jsonify({"profile": profile})
        return jsonify({"error": "Could not parse AI response into profile JSON"}), 500
    except Exception as e:
        return jsonify({"error": f"Parse error: {str(e)}", "raw": result}), 500

def is_ai_role(jd, role_type):
    ai_terms = ["ai", "artificial intelligence", "machine learning", "ml", "llm",
                "generative ai", "genai", "nlp", "gpt", "claude", "openai",
                "foundation model", "large language model", "ai product", "data science"]
    text = (jd + " " + role_type).lower()
    return any(t in text for t in ai_terms)


# ─── ROUTES ───────────────────────────────────────────────

@app.route("/", methods=["GET", "HEAD"])
@app.route("/index.html")
def index():
    try:
        return render_template("index.html")
    except Exception as e:
        app.logger.error(f"Template error: {e}")
        return f"<h2>App is running!</h2><p>Template error: {e}</p><p>BASE_DIR: {BASE_DIR}</p>", 500

def _inject_ai_projects(resume_text):
    """Always inject exactly ONE AI & Personal Projects section before ACADEMIC QUALIFICATION.
    Strips any AI-generated version first, then inserts our fixed canonical block."""
    import re as _re

    PROJECTS = (
        "\nAI & PERSONAL PROJECTS:\n"
        "AI Trade Analysis Platform & Job Hunt Automation App | Python, Flask, Claude API, Supabase, Render | 2025\n"
        "- Independently designed and deployed two live AI-powered applications: a stock monitoring platform "
        "(https://stock-monitor-8ak6.onrender.com) using NLP and machine learning to monitor stock prices and provide "
        "real-time trade insights, and a job-hunt automation app (https://job-hunt-app-r7my.onrender.com) that generates "
        "tailored ATS-optimised resumes and cover letters, scores job fit via AI ranking, and tracks applications via a Kanban board\n"
        "- Gained hands-on experience across the full product lifecycle: problem definition, prompt engineering, "
        "LLM API integration (Claude API), Flask backend, Supabase database, LinkedIn bookmarklet development, "
        "and end-to-end Render cloud deployment — independently shipped both products from scratch"
    )

    # Strip ALL existing AI projects blocks (AI-generated or previously injected)
    resume_text = _re.sub(
        r'\n+AI & PERSONAL PROJECTS:.*?(?=\nACADEMIC QUALIFICATION|\nEDUCATION|\Z)',
        '', resume_text, flags=_re.DOTALL | _re.IGNORECASE
    )
    resume_text = _re.sub(
        r'\n+PERSONAL PROJECTS:.*?(?=\nACADEMIC QUALIFICATION|\nEDUCATION|\Z)',
        '', resume_text, flags=_re.DOTALL | _re.IGNORECASE
    )
    resume_text = _re.sub(
        r'\n+AI PROJECTS:.*?(?=\nACADEMIC QUALIFICATION|\nEDUCATION|\Z)',
        '', resume_text, flags=_re.DOTALL | _re.IGNORECASE
    )

    # Insert before ACADEMIC QUALIFICATION or EDUCATION section
    for marker in ['ACADEMIC QUALIFICATION:', 'ACADEMIC QUALIFICATION', 'EDUCATION & CERTIFICATIONS:', 'EDUCATION:']:
        idx = resume_text.find(marker)
        if idx != -1:
            resume_text = resume_text[:idx].rstrip() + PROJECTS + '\n\n' + resume_text[idx:]
            return resume_text

    # Fallback: append at end
    return resume_text.rstrip() + PROJECTS


@app.route("/api/tailor-resume", methods=["POST"])
def tailor_resume():
    data = request.json
    jd = data.get("jd", "")
    role_type = data.get("roleType", "Business Analyst")
    matched_keywords = data.get("matchedKeywords", [])   # from rank_jobs scoring
    ai_role = is_ai_role(jd, role_type)
    P = get_active_profile()
    framing = build_product_framing(P)

    # Build keyword line for prompt — use pre-extracted keywords if available
    if matched_keywords:
        kw_line = f"PRE-EXTRACTED JD KEYWORDS (already matched against candidate profile — use ALL of these verbatim in resume): {', '.join(matched_keywords)}"
    else:
        kw_line = "Extract keywords from the JD above and weave them throughout the resume."

    # Build structured profile text (easier for LLM than JSON)
    exp_text = ""
    for exp in P.get('experience', []):
        exp_text += f"\n--- {exp.get('role','')} at {exp.get('company','')} ({exp.get('period','')}) ---\n"
        for b in exp.get('bullets', []):
            exp_text += f"  - {b}\n"
        for a in exp.get('achievements', []):
            exp_text += f"  Achievement: {a}\n"

    skills_text = ', '.join(P.get('skills', []))
    edu_text = '\n'.join([f"  - {e.get('degree','')} — {e.get('school','')} ({e.get('period','')})" for e in P.get('education', [])])
    proj_text = ""
    for proj in P.get('projects', []):
        proj_text += f"\n  Project: {proj.get('title','')} ({proj.get('period','')}) — {proj.get('tech','')}\n"
        proj_text += f"  URL: {proj.get('url','')}\n"
        for b in proj.get('bullets', []):
            proj_text += f"  - {b}\n"

    num_exp = len(P.get('experience', []))

    prompt = f"""You are an expert ATS resume writer. Write a COMPLETE 2-page resume for {P.get('name','')} targeting: {role_type} at a company matching this JD.

JOB DESCRIPTION:
{jd}

{kw_line}

===== CANDIDATE MASTER DATA (use ALL of this) =====
Name: {P.get('name','')}
Phone: {P.get('mobile','')}
Email: {P.get('email','')}
LinkedIn: {P.get('linkedin','')}
Address: {P.get('address','')}
Certification: SAFe 6.0 Product Owner/Product Management

HEADLINE OPTIONS (pick most relevant to JD):
- Business Analyst Lead | Product Owner | Sales & Operations Excellence
- Lead Business Analyst / Enterprise Product Owner
- Digital Product Manager | Product Owner (AI & Data Platforms)
- Business Analyst Lead | Product Owner | Product Manager

SUMMARY OPTIONS (adapt and blend most relevant to JD):
Option A: Lead Business Analyst with 5+ years of experience supporting senior leadership through business planning, financial operations, and cross-functional project execution. Proven ability to streamline processes, manage executive stakeholder communications, track KPIs, drive operational alignment and cost profitability.
Option B: Lead Business Analyst / Enterprise Product Owner with 5+ years of experience delivering large-scale enterprise platforms across financial services. Proven expertise in owning product roadmaps, custom-built systems, translating business needs into scalable technical solutions, and driving platform adoption across multiple business units.
Option C: Digital Product Manager / Product Owner with 5+ years of experience delivering large-scale digital platforms across financial services and enterprise environments. Proven expertise in owning product roadmaps, managing Agile delivery, translating user needs into scalable solutions, and driving data-informed product decisions.
Option D (Fintech/Revolut-style): Results-driven Product Owner with 5+ years in fintech and banking, skilled in SEO, GEO, and CRO, with expertise in driving website optimisation and owning product roadmaps. Proficient in SAFe 6.0 and Agile, with strong understanding of design best practices and user experience.

SKILLS POOL (use all relevant + add JD keywords):
Data visualization tools: Tableau, Power BI
Programming: PSQL, Python basics
Key Modules in Masters: Financial Planning and Analysis, Operations management
Others: Management Consulting, Agile, SAFe 6.0, JIRA, Excel, Microsoft Project, Product Vision and road mapping, Business Analysis, Risk Mitigation & Change Management, Budget Forecasting & Variance Analysis, KPI Tracking & Dashboard Reporting, Event & Workshop Facilitation, AI-enabled Product Integration, MVP Definition & Go-To-Market, SEO, GEO, CRO, API integrations, Generative AI, LLM, Prompt Engineering, Stakeholder Management

EXPERIENCE — USE ALL BULLETS BELOW, SELECT MOST JD-RELEVANT ONES:

KPMG, Singapore | Feb 2021 – Present | Lead Business Analyst / Functional Consultant / Digital Product Manager
Available bullets (pick 6-8 most relevant to JD):
- Defined and drove product vision, roadmap, and delivery strategy for large-scale digital transformation initiatives for financial institutions
- Owned and prioritised product backlog, ensuring alignment with business objectives, regulatory requirements and customer experience goals
- Partnered with Enterprise Singapore / GovTech on large-scale digital transformation projects
- Supported executive decision-making through As-Is/To-Be analysis and streamlined documentation of strategic processes across finance teams
- Collaborated with cross-regional stakeholders (UX, architecture, and finance) to drive alignment across engineering and operations functions
- Managed internal reporting cycles, change requests, and approval workflows contributing to a 5% increase in project profitability
- Drove product scope decisions through impact analysis, generating ~5% additional business value
- Led sprint ceremonies including Planning, Reviews, Retros, and PI Planning across a multi-squad programme, resulting in 30% increase in team velocity
- Managed 3rd party vendors, conducted go-live planning, and led data migrations from legacy systems ensuring 99.9% uptime
- Designed and executed end-to-end test scenarios on Loan IQ applications resulting in 25% reduction in defects
- Trained vendors and managed go-live execution plans, ensuring process continuity across business units
- Conducted internal workshops to drive operational readiness and team-wide coordination
- Analysed complex problems, root cause analysis and creating innovative solutions with UX, architecture and software solutioning teams. Experience in API integrations
- Design, document and execute end-to-end test scenarios on Loan IQ Product Applications including M&A, Trade, WCL, FA and raise defects using JIRA
- Expertise in Monthly Reporting Reconciliation, Exposure Calculation, Commitment & Credit Limit Calculation for Banks, Financial Loan Products, Lending Capital Loans
- Managing 3rd party vendors and Client Senior Stakeholders, train end users, conducting workshops, Data Migrations from legacy systems
- Supported sales and distribution process efficiency by streamlining reporting, reducing cycle times, improving customer-facing loan/trade workflows
- Leading Testers, Interns and Junior Business Analysts for Scrum activities: Planning, Closure and Retro
- Supported automation and AI-enabled solution design including API integrations, improving efficiency and reducing manual effort by 30 man-days
- Built business cases and executive presentations for new product features, securing stakeholder approval
- Drove data-informed product decisions by defining KPIs, analysing performance metrics and identifying product improvement opportunities
- Worked on E-commerce website analysing customer engagement patterns leading to 5% increase in customer retention
Key Achievements:
- Performed accurate Impact Analysis and persuaded clients to approve Change request items, generating additional profit of ~5% of Project Cost
- Analysed and provided solutioning for Automated Interest computation (financial services), saving 30 man-days of work
- Led team for interim in critical phase during end of Sprints and beginning of SIT, maintaining project delivery timeline

J.P. Morgan | Oct 2023 – Jan 2024 | Asset Management Virtual Internship
Available bullets (pick 3-4 most relevant):
- Gathered product requirements from trading/execution teams to build robust investor profiles using UX/UI design principles
- Helped traders and clients onboard, build stronger investment portfolios and offered market-leading investment solutions
- Emailed questionnaires to clients to build robust investor profiles and extract essential information such as investment objectives and risk appetite
- Performed quantitative fundamental analysis of 5 stocks, recommended to 2 clients based on risk metrics resulting in 10% increase in portfolio value
- Measured portfolio performance via KPIs: Annual Portfolio Return, Portfolio Variance, Standard Deviation

Amazon Inc, India | Mar 2018 – Mar 2019 | Business Analyst
Available bullets (pick 2-3 most relevant):
- Built real-time quality monitoring dashboards using Power BI from SQL Server and MS Excel, resulting in 20% reduction in quality issues
- Translated business requirements into functional and non-functional specifications ensuring 95% stakeholder satisfaction
- Worked closely with stakeholders to understand needs, scope problems and develop business cases on turning data into actionable information
- Analysed and visualised operational data using Tableau and Power BI, resulting in 15% increase in operational efficiency

EDUCATION:
Master of Science Engineering Business Management, Coventry University, UK (Jul 2019 – Nov 2020)
Bachelor of Engineering Electronics & Communication Engineering, Anna University, India (Jul 2012 – Jun 2016)
SAFe 6.0 certified Product Owner/Product Management
===================================================

OUTPUT FORMAT RULES — FOLLOW EXACTLY:
1. First line: Amretha Karthikeyan  (just the name, nothing else)
2. Line 2: #02-321 153 Gangsa Road, Singapore-670153
3. Line 3: Mobile: +65-90256503, email: amretha.ammu@gmail.com
4. Line 4: https://www.linkedin.com/in/amretha-nishanth-534b39101/
5. Line 5: [Most relevant headline from options above]
6. Blank line
7. PROFESSIONAL SUMMARY:
8. [4 sentences, blend from options above, weave in JD keywords]
9. Blank line
10. SKILL SET:
11. Data visualization tools: Tableau, Power BI[+ JD tools]
12. Programming: PSQL, Python basics[+ JD languages]
13. Others: [all relevant skills + JD keywords, comma separated]
14. Certification: Scaled Agile Framework 6.0 Product Owner/Product Management
15. Blank line
16. PROFESSIONAL EXPERIENCE:
17. KPMG, Singapore  Feb 2021 – Present
18. [Most relevant job title — no slashes at start/end]
19. - [bullet 1]
20. - [bullet 2] ... (6-8 bullets)
21. Key Achievements:
22. - [achievement 1]
23. - [achievement 2]
24. - [achievement 3]
25. Blank line
26. J.P. Morgan  Oct 2023 – Jan 2024
27. Asset Management Virtual Internship
28. - [3-4 bullets]
29. Blank line
30. Amazon Inc, India  Mar 2018 – Mar 2019
31. Business Analyst
32. - [2-3 bullets]
33. Blank line
34. ACADEMIC QUALIFICATION:
35. Master of Science Engineering Business Management  Jul 2019 – Nov 2020
36. Coventry University, UK
37. Bachelor of Engineering  Jul 2012 – Jun 2016
38. Electronics & Communication Engineering, Anna University, India
39. Certification: Scaled Agile Framework 6.0 Product Owner/Product Management

CRITICAL RULES:
- Plain text ONLY. No **, no #, no ```, no markdown
- Section headers in ALL CAPS exactly as shown
- Bullets use "- " prefix
- Job titles on their own line, no slashes before or after
- Target 700-850 words total (fills 2 pages properly)
- Weave in at least 12 exact keyword phrases from the JD
- Do NOT write "HEADER" anywhere — start with the candidate name
- Do NOT mention the target company name (e.g. Revolut, any specific company) anywhere in the resume — language must be generic and transferable
- Do NOT add any closing paragraphs, cover-letter-style text, or "overall I believe..." summaries after the ACADEMIC QUALIFICATION section"""

    result = call_claude(prompt, max_tokens=8192)
    result = _inject_ai_projects(result)
    return jsonify({"result": result, "isAiRole": ai_role})

@app.route("/api/cover-letter", methods=["POST"])
def cover_letter():
    data = request.json
    jd = data.get("jd", "")
    role_type = data.get("roleType", "Business Analyst")
    company = data.get("company", "the company")
    ai_role = is_ai_role(jd, role_type)

    P = get_active_profile()
    framing = build_product_framing(P)
    # Build achievements from profile experience
    achievements_text = ""
    for exp in P.get('experience', []):
        for ach in exp.get('achievements', []):
            achievements_text += f"- {ach}\n"
    if P.get('certification'):
        achievements_text += f"- Certified: {P['certification']}\n"
    if P.get('aiProjectUrl'):
        achievements_text += f"- Personal Project: {P.get('aiProjectUrl','')}\n"

    prompt = f"""Write a professional 300-350 word cover letter for {P['name']} applying to {role_type} at {company}.
{framing}

KEY ACHIEVEMENTS:
{achievements_text}

JOB DESCRIPTION:
{jd}

{"IMPORTANT — AI ROLE: Mention the project at " + P.get('aiProjectUrl','') + " as proof of hands-on AI product development. Include the URL." if ai_role else ""}

ATS & RECRUITER OPTIMISATION:
1. Mirror the EXACT job title and 5-8 key phrases from the JD in the letter.
2. Use confident product language, not consulting jargon.
3. Include specific metrics (5% value, 30 man-days) for credibility.
4. Reference the company name and role title at least twice.
5. Keep paragraphs short (3-4 sentences max) for easy scanning.

Write a compelling cover letter that:
1. Opens with a confident hook referencing the specific role and company, positioning as a product builder not a service provider
2. Highlights KPMG metrics (5% value, 30 man-days) in context of what JD requires
3. {"Mentions live AI project with URL as key differentiator" if ai_role else "Bridges consulting delivery to product ownership with specific JD alignment"}
4. Shows genuine, specific enthusiasm for {company} — reference what they do
5. Ends with a clear, action-oriented call to action

Exactly 300-350 words. No consulting jargon. Sound like a product person. Weave JD keywords naturally throughout."""

    result = call_claude(prompt)
    return jsonify({"result": result})

@app.route("/api/interview-prep", methods=["POST"])
def interview_prep():
    data = request.json
    company = data.get("company", "the company")
    role_type = data.get("roleType", "Business Analyst")
    jd = data.get("jd", "")

    P = get_active_profile()
    framing = build_product_framing(P)
    # Build candidate summary from profile
    exp_lines = ""
    for exp in P.get('experience', []):
        bullets_preview = '; '.join(exp.get('bullets', [])[:2])
        exp_lines += f"- {exp.get('company','')} ({exp.get('period','')}): {exp.get('role','')}. {bullets_preview}\n"
    skills_str = ', '.join(P.get('skills', [])[:12])
    proj_url = P.get('aiProjectUrl', '')

    prompt = f"""Generate a comprehensive interview prep guide for {P['name']} interviewing at {company} for {role_type}.
{framing}

CANDIDATE:
{exp_lines}
- Certified: {P.get('certification','')}
- Skills: {skills_str}
{('- Project: ' + proj_url) if proj_url else ''}
{"JD: " + jd if jd else ""}

Create prep with these EXACT sections:

## 5 Behavioral Questions with STAR Answers
For each: the question, then full STAR answer using the candidate's real experience with specific metrics.

## 5 Technical Questions for {role_type}
Questions with model answers specific to this role.

## 3 Things to Research About {company}
Specific actionable research areas.

## 5 Smart Questions to Ask the Interviewer
Product-minded questions that signal ownership thinking.

## Salary Negotiation Tip
Specific tip based on the candidate's certifications and experience level."""

    result = call_claude(prompt)
    return jsonify({"result": result})


# ─── INTERACTIVE AI INTERVIEW COACH ─────────────────────────────────────────

# In-memory session store (per-server; for multi-server use Redis/Supabase)
_interview_sessions = {}

def _scrape_company_intel(company):
    """Try to gather company interview intelligence from public sources."""
    intel = {"glassdoor": None, "general": None}
    try:
        import requests as http_req
        # Try Glassdoor-style search via Google
        headers = {"User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36"}
        queries = [
            f"{company} interview questions glassdoor",
            f"{company} interview process experience",
        ]
        snippets = []
        for q in queries[:1]:  # limit to 1 query to be fast
            try:
                url = f"https://www.google.com/search?q={q.replace(' ', '+')}&num=5"
                r = http_req.get(url, headers=headers, timeout=8)
                from bs4 import BeautifulSoup
                soup = BeautifulSoup(r.text, "html.parser")
                for div in soup.select(".BNeawe.s3v9rd"):
                    text = div.get_text(strip=True)
                    if len(text) > 40:
                        snippets.append(text[:300])
                    if len(snippets) >= 5:
                        break
            except Exception:
                pass
        if snippets:
            intel["glassdoor"] = snippets
    except Exception:
        pass
    return intel


@app.route("/api/interview/start", methods=["POST"])
def interview_start():
    """Start an interactive interview session."""
    data = request.json or {}
    role = data.get("role", "").strip()
    company = data.get("company", "").strip()
    interview_type = data.get("type", "behavioral")
    jd = data.get("jd", "").strip()
    resume_text = data.get("resume", "").strip()

    if not role or not company:
        return jsonify({"error": "Role and company are required"}), 400

    P = get_active_profile()

    # Build candidate context from profile + optional resume
    candidate_info = f"Name: {P.get('name', 'Candidate')}\n"
    if P.get('headline'):
        candidate_info += f"Headline: {P['headline']}\n"
    if P.get('summary'):
        candidate_info += f"Summary: {P['summary'][:300]}\n"
    for exp in P.get('experience', [])[:3]:
        bullets = '; '.join(exp.get('bullets', [])[:2])
        candidate_info += f"- {exp.get('company','')} | {exp.get('role','')} ({exp.get('period','')}): {bullets}\n"
    if resume_text:
        candidate_info += f"\nResume excerpt:\n{resume_text[:800]}\n"

    # Type-specific instruction
    type_instructions = {
        "behavioral": """Focus on behavioral and situational questions using the STAR method.
Ask questions like "Tell me about a time when..." and "How would you handle...".
Cover: leadership, conflict, failure, teamwork, prioritization, stakeholder management.
After each answer, evaluate their STAR structure and probe deeper.""",

        "technical": """Focus on technical and case study questions relevant to the role.
Include: system design, data analysis, technical problem-solving, SQL/analytics scenarios.
Ask follow-up questions to test depth of knowledge.""",

        "product": f"""Focus on product sense, strategy, and case study questions.
Use frameworks from top PM prep (Exponent, PrepLounge style).
Include: product design, metrics/KPIs, go-to-market, prioritization frameworks, estimation.
After each answer, evaluate their structured thinking.""",

        "mixed": """Conduct a realistic full mock interview mixing:
- 2 behavioral/situational questions (STAR method)
- 2 technical/role-specific questions
- 1 product sense or case question
Transition naturally between types like a real interviewer.""",
    }

    # Try to get company intel
    company_intel = _scrape_company_intel(company)
    intel_context = ""
    if company_intel.get("glassdoor"):
        intel_context = f"\nCompany interview intelligence (from web research):\n" + "\n".join(f"- {s}" for s in company_intel["glassdoor"][:3])

    session_id = f"session_{id(data)}_{__import__('time').time()}"

    system_prompt = f"""You are an expert AI interview coach conducting a realistic mock interview.
You are interviewing a candidate for the role of **{role}** at **{company}**.

CANDIDATE PROFILE:
{candidate_info}

{"JOB DESCRIPTION:\n" + jd[:1000] if jd else ""}
{intel_context}

INTERVIEW TYPE: {interview_type}
{type_instructions.get(interview_type, type_instructions['behavioral'])}

RULES:
1. Ask ONE question at a time. Wait for the candidate's answer before continuing.
2. After each answer, provide:
   - A brief score (1-10) with reasoning
   - Specific feedback on what was good and what to improve
   - A follow-up or new question
3. Be encouraging but honest. Point out weak areas constructively.
4. Track which competencies you've covered.
5. If the answer is vague, probe deeper — "Can you be more specific?" or "What was the measurable outcome?"
6. Reference the candidate's actual experience from their profile when asking questions.
7. Vary difficulty — start with a warm-up question, then increase complexity.
8. Format your response as:
   **Score: X/10** [brief reason]
   **Feedback:** [specific feedback]
   **Next Question:** [the next question]
   (For the FIRST message, skip score/feedback and just ask an opening question with a brief welcome.)"""

    # Generate the first question
    first_msg_prompt = f"""{system_prompt}

Start the interview now. Welcome the candidate warmly, mention the role and company, and ask your first question.
Keep the welcome to 2 sentences max, then ask the question."""

    first_response = call_claude(first_msg_prompt)

    # Store session
    _interview_sessions[session_id] = {
        "system_prompt": system_prompt,
        "messages": [
            {"role": "assistant", "content": first_response}
        ],
        "role": role,
        "company": company,
        "type": interview_type,
        "scores": [],
        "started_at": __import__('time').time(),
        "company_intel": company_intel,
    }

    return jsonify({
        "session_id": session_id,
        "message": first_response,
        "company_intel": company_intel,
    })


@app.route("/api/interview/respond", methods=["POST"])
def interview_respond():
    """Process a candidate's answer and generate AI follow-up."""
    data = request.json or {}
    session_id = data.get("session_id", "")
    answer = data.get("answer", "").strip()

    if not session_id or session_id not in _interview_sessions:
        return jsonify({"error": "Invalid or expired session"}), 400
    if not answer:
        return jsonify({"error": "Please provide an answer"}), 400

    session = _interview_sessions[session_id]
    session["messages"].append({"role": "user", "content": answer})

    # Build conversation for Claude
    conversation = session["system_prompt"] + "\n\n"
    conversation += "CONVERSATION SO FAR:\n"
    for msg in session["messages"]:
        prefix = "INTERVIEWER" if msg["role"] == "assistant" else "CANDIDATE"
        conversation += f"\n{prefix}: {msg['content']}\n"

    conversation += "\nINTERVIEWER (now respond with score, feedback, and next question):"

    response = call_claude(conversation)

    session["messages"].append({"role": "assistant", "content": response})

    # Extract score if present
    import re
    score_match = re.search(r'\*?\*?Score:\s*(\d+)/10', response)
    if score_match:
        session["scores"].append(int(score_match.group(1)))

    return jsonify({
        "message": response,
        "question_count": sum(1 for m in session["messages"] if m["role"] == "assistant"),
        "answer_count": sum(1 for m in session["messages"] if m["role"] == "user"),
        "avg_score": round(sum(session["scores"]) / len(session["scores"]), 1) if session["scores"] else None,
    })


@app.route("/api/interview/end", methods=["POST"])
def interview_end():
    """End interview session and generate comprehensive summary."""
    data = request.json or {}
    session_id = data.get("session_id", "")

    if not session_id or session_id not in _interview_sessions:
        return jsonify({"error": "Invalid or expired session"}), 400

    session = _interview_sessions[session_id]

    # Build full transcript
    transcript = ""
    for msg in session["messages"]:
        prefix = "🤖 Interviewer" if msg["role"] == "assistant" else "👤 You"
        transcript += f"\n{prefix}:\n{msg['content']}\n"

    summary_prompt = f"""You conducted a mock interview for {session['role']} at {session['company']}.
Type: {session['type']}

Full transcript:
{transcript}

Provide a comprehensive session summary with:

## Overall Performance Score
Give an overall score out of 10 with detailed reasoning.

## Strengths Demonstrated
List 3-5 specific strengths shown during the interview, with examples from their answers.

## Areas for Improvement
List 3-5 specific areas to improve, with actionable recommendations.

## STAR Method Assessment
Rate their use of the STAR method (Situation, Task, Action, Result) in behavioral answers. Which component was weakest?

## Communication Analysis
Assess: clarity, conciseness, confidence, structure, use of metrics/data.

## Key Recommendations
Top 3 actionable things to practice before the real interview.

## Sample Improved Answer
Take their weakest answer and rewrite it as an ideal response.

Be specific, reference their actual answers, and be constructive."""

    summary = call_claude(summary_prompt)

    # Clean up session
    result = {
        "summary": summary,
        "transcript": transcript,
        "total_questions": sum(1 for m in session["messages"] if m["role"] == "assistant"),
        "total_answers": sum(1 for m in session["messages"] if m["role"] == "user"),
        "avg_score": round(sum(session["scores"]) / len(session["scores"]), 1) if session["scores"] else None,
        "duration_seconds": int(__import__('time').time() - session["started_at"]),
    }

    del _interview_sessions[session_id]
    return jsonify(result)


@app.route("/api/interview/company-intel", methods=["POST"])
def interview_company_intel():
    """Fetch company interview intelligence."""
    data = request.json or {}
    company = data.get("company", "").strip()
    role = data.get("role", "").strip()

    if not company:
        return jsonify({"error": "Company name required"}), 400

    intel = _scrape_company_intel(company)

    # Also ask AI for company-specific insights
    prompt = f"""Provide interview intelligence for {role or 'a candidate'} interviewing at {company}:

## Company Overview
Brief company description, culture, and values (2-3 sentences).

## Interview Process
Typical interview stages and what to expect at {company}.

## Common Interview Questions at {company}
List 5 questions commonly asked at {company} based on known patterns.

## Company-Specific Tips
3 tips specifically for succeeding at a {company} interview.

## Key Values & Culture Fit
What {company} looks for in candidates — culture signals to demonstrate.

## Recent News & Talking Points
2-3 recent developments at {company} worth mentioning in the interview.

Be specific to {company}. If you don't have specific info, provide educated guidance based on the company's industry and size."""

    ai_intel = call_claude(prompt)

    return jsonify({
        "ai_intel": ai_intel,
        "web_snippets": intel.get("glassdoor", []),
    })


@app.route("/api/full-kit", methods=["POST"])
def full_kit():
    data = request.json
    company = data.get("company", "")
    role = data.get("role", "")
    role_type = data.get("roleType", "Business Analyst")
    jd = data.get("jd", "")
    ai_role = is_ai_role(jd, role_type)

    P = get_active_profile()
    framing = build_product_framing(P)
    profile_str = json.dumps({k: v for k, v in P.items()}, indent=2)
    proj_url = P.get('aiProjectUrl', '')

    resume_prompt = f"Write ATS-optimised resume for {P['name']} applying to {role} at {company} ({role_type}). {framing} Profile: {profile_str}. JD: {jd}. {'AI role: feature project ' + proj_url + ' prominently.' if ai_role else ''} ATS rules: mirror exact JD keywords, use standard section headers (Professional Summary, Core Skills, Professional Experience, Education), include metrics in every bullet, single-column format, no tables."
    cover_prompt = f"Write 300-word cover letter for {P['name']} for {role} at {company}. Profile: {profile_str}. {'Mention project: ' + proj_url if ai_role else ''} Mirror key phrases from JD: {jd[:500]}. Product language, no consulting jargon. Reference company name and role at least twice."
    prep_prompt = f"Give top 5 interview questions for {role_type} at {company} with brief model answers for {P['name']}. Profile summary: {P.get('summary','')}. JD context: {jd[:500]}. Include STAR-format answers with real metrics."

    import concurrent.futures
    with concurrent.futures.ThreadPoolExecutor() as executor:
        r_future = executor.submit(call_claude, resume_prompt)
        c_future = executor.submit(call_claude, cover_prompt)
        p_future = executor.submit(call_claude, prep_prompt)
        resume = r_future.result()
        cover = c_future.result()
        prep = p_future.result()

    return jsonify({"resume": resume, "cover": cover, "prep": prep, "isAiRole": ai_role})

@app.route("/api/follow-up", methods=["POST"])
def follow_up():
    data = request.json
    company = data.get("company", "the company")
    role = data.get("role", "the role")
    days = data.get("days", 7)

    prompt = f"""Write a polite 3-line follow-up email from Amretha Karthikeyan about her application for {role} at {company}, submitted {days} days ago.
Include: subject line, brief message referencing the role, continued interest, offer to provide more info.
Under 80 words. Ready to copy-paste. Professional and confident."""

    result = call_claude(prompt)
    return jsonify({"result": result})

@app.route("/api/speed-kit", methods=["POST"])
def speed_kit():
    data = request.json
    company = data.get("company", "this company")
    role = data.get("role", "this role")

    prompt = f"""Write a genuine 3-sentence "Why do you want to work at {company}?" answer for Amretha Karthikeyan, a SAFe 6.0 PO/Lead BA transitioning from KPMG to an in-house {role} role. Be specific to {company}'s product/market. Sound like a product person who wants to build. No consulting language."""

    result = call_claude(prompt)
    return jsonify({"result": result})


@app.route("/api/generic", methods=["POST"])
def generic():
    data = request.json
    prompt = data.get("prompt", "")
    system = data.get("systemPrompt", "")
    full_prompt = f"{system}\n\n{prompt}" if system else prompt
    result = call_claude(full_prompt)
    return jsonify({"result": result})


@app.route("/api/health", methods=["GET"])
def health_check():
    """Debug endpoint to check environment configuration."""
    sb = get_supabase()
    sb_status = "connected" if sb else "NOT configured"
    if sb:
        try:
            res = sb.table("jobs").select("id", count="exact").execute()
            sb_status = f"connected ({res.count or 0} jobs in DB)"
        except Exception as e:
            sb_status = f"connected but query failed: {str(e)[:100]}"
    return jsonify({
        "status": "ok",
        "groq_api_key": "SET" if GROQ_API_KEY else "MISSING",
        "supabase_url": "SET" if SUPABASE_URL else "MISSING",
        "supabase_key": "SET" if SUPABASE_KEY else "MISSING",
        "supabase_key_length": len(SUPABASE_KEY) if SUPABASE_KEY else 0,
        "supabase_key_prefix": (SUPABASE_KEY[:20] + "...") if SUPABASE_KEY else "empty",
        "supabase_status": sb_status,
        "supabase_error": _supabase_error,
        "supabase_url_preview": (SUPABASE_URL[:40] + "...") if SUPABASE_URL else "empty"
    })


@app.route("/api/jobs", methods=["GET"])
def get_jobs():
    """Load all jobs from Supabase."""
    sb = get_supabase()
    if not sb:
        return jsonify({"error": "Supabase not configured", "jobs": []}), 200
    try:
        res = sb.table("jobs").select("*").order("created_at", desc=False).execute()
        return jsonify({"jobs": res.data or []})
    except Exception as e:
        return jsonify({"error": str(e), "jobs": []}), 200


@app.route("/api/jobs/upsert", methods=["POST"])
def upsert_jobs():
    """Save/update jobs to Supabase. Upserts by job id."""
    sb = get_supabase()
    if not sb:
        return jsonify({"error": "Supabase not configured"}), 200
    data = request.json
    jobs = data.get("jobs", [])
    if not jobs:
        return jsonify({"ok": True, "count": 0})
    try:
        # Whitelist fields we want to persist
        # Only include resume/cover docx if they're actually sent (to avoid
        # overwriting existing docs when the frontend syncs lightweight payloads)
        def clean(j):
            row = {
                "id":               str(j.get("id", "")),
                "role":             j.get("role", ""),
                "company":          j.get("company", ""),
                "status":           j.get("status", "saved"),
                "url":              j.get("url", ""),
                "linkedInId":       j.get("linkedInId", ""),
                "jd":               (j.get("jd") or "")[:8000],
                "roleType":         j.get("roleType", ""),
                "source":           j.get("source", ""),
                "salary":           j.get("salary", ""),
                "dateApplied":      j.get("dateApplied", ""),
                "aiScore":          j.get("aiScore"),
                "aiLabel":          j.get("aiLabel", ""),
                "aiReason":         j.get("aiReason", ""),
                "aiPriority":       j.get("aiPriority", ""),
                "notes":            j.get("notes", ""),
                "resume_variant":   j.get("resume_variant", ""),
                "resume_filename":  j.get("resume_filename", ""),
                "cover_filename":   j.get("cover_filename", ""),
                "resume_generated_at": j.get("resume_generated_at", ""),
            }
            # Only include large binary fields if they're explicitly present & non-empty
            # This prevents the frontend lightweight sync from clearing doc data
            if j.get("resume_docx_b64"):
                row["resume_docx_b64"] = j["resume_docx_b64"][:500000]
            if j.get("cover_docx_b64"):
                row["cover_docx_b64"] = j["cover_docx_b64"][:500000]
            return row
        cleaned = [clean(j) for j in jobs if j.get("id")]
        # Batch upsert for reliability
        BATCH = 30
        total = 0
        for i in range(0, len(cleaned), BATCH):
            batch = cleaned[i:i+BATCH]
            sb.table("jobs").upsert(batch, on_conflict="id").execute()
            total += len(batch)
        return jsonify({"ok": True, "count": total})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/jobs/delete", methods=["POST"])
def delete_job():
    """Delete a job from Supabase by id."""
    sb = get_supabase()
    if not sb:
        return jsonify({"error": "Supabase not configured"}), 200
    data = request.json
    job_id = data.get("id")
    if not job_id:
        return jsonify({"error": "No id"}), 400
    try:
        sb.table("jobs").delete().eq("id", job_id).execute()
        return jsonify({"ok": True})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/jobs/clear-all", methods=["POST"])
def clear_all_jobs():
    """Delete every job from Supabase — fresh start."""
    sb = get_supabase()
    if not sb:
        return jsonify({"error": "Supabase not configured"}), 200
    try:
        # Delete all rows — Supabase requires a filter, use neq on a always-true condition
        sb.table("jobs").delete().neq("id", "___never___").execute()
        return jsonify({"ok": True})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


def _create_docx_from_text(text, title="Document"):
    """
    Render AI-generated resume text into a formatted .docx matching Amretha CV style.
    Rules:
      - First non-empty line = candidate name (large, centered, bold)
      - Lines 2-5 with contact info = centered, bold
      - Section headers (ALL CAPS known keywords) = bold, bottom-border, no bullets
      - Company + date lines = bold, no bullets
      - Job titles (role keywords, no year) = bold, no bullets, no slashes
      - Skill category lines "Label: values" = label bold, values normal
      - Education body text = not bold
      - Bullet lines starting with "- " = proper list bullets
      - Everything else = normal body text
    """
    from docx import Document as DocxDocument
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import io, re

    FONT        = 'Times New Roman'
    BODY_SIZE   = Pt(11)
    HDR_SIZE    = Pt(12)
    NAME_SIZE   = Pt(16)

    SECTION_HEADERS = {
        'PROFESSIONAL SUMMARY','SUMMARY','CORE SKILLS','SKILL SET','SKILLS',
        'PROFESSIONAL EXPERIENCE','EXPERIENCE','EDUCATION & CERTIFICATIONS',
        'ACADEMIC QUALIFICATION','EDUCATION','CERTIFICATIONS','PROJECTS',
        'AI & INNOVATION','KEY ACHIEVEMENTS','QUALIFICATIONS','CONTACT',
        'PRODUCT IMPACT','INTEREST AREAS','AI & PERSONAL PROJECTS',
        'PERSONAL PROJECTS','AI PROJECTS',
    }
    EDUCATION_SECTIONS = {
        'EDUCATION & CERTIFICATIONS','ACADEMIC QUALIFICATION','EDUCATION','CERTIFICATIONS'
    }
    SKILL_CATS = [
        'Data visualization tools','Data Visualization','Programming',
        'Key Modules in Masters','Others','Certification','Certifications',
        'Methodologies','Tools','Soft Skills','Domain','Analytics',
        'Project Management','Languages','AI & Automation',
    ]
    JOB_TITLE_KW = [
        'analyst','manager','owner','lead','consultant','engineer','director',
        'associate','intern','officer','specialist','coordinator','head',
        'senior','junior','internship','product','digital',
    ]

    def add_border(p):
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bot  = OxmlElement('w:bottom')
        bot.set(qn('w:val'),   'single')
        bot.set(qn('w:sz'),    '6')
        bot.set(qn('w:space'), '1')
        bot.set(qn('w:color'), '000000')
        pBdr.append(bot)
        pPr.append(pBdr)

    def rf(run, size=None, bold=False):
        run.font.name = FONT
        run.font.size = size or BODY_SIZE
        run.bold      = bold

    def plain_para(doc, text_str, bold=False, before=0, after=3):
        p   = doc.add_paragraph()
        run = p.add_run(text_str)
        rf(run, bold=bold)
        p.paragraph_format.space_before = Pt(before)
        p.paragraph_format.space_after  = Pt(after)
        return p

    def clean_title(s):
        return re.sub(r'^[\s/\-–|]+|[\s/\-–|]+$', '', s).strip()

    doc = DocxDocument()
    doc.styles['Normal'].font.name = FONT
    doc.styles['Normal'].font.size = BODY_SIZE
    for sec in doc.sections:
        sec.top_margin    = Inches(0.6)
        sec.bottom_margin = Inches(0.6)
        sec.left_margin   = Inches(0.75)
        sec.right_margin  = Inches(0.75)
    try:
        doc.styles['List Bullet'].font.name = FONT
        doc.styles['List Bullet'].font.size = BODY_SIZE
    except Exception:
        pass

    lines        = [l for l in text.split('\n')]

    # ── HARD TRUNCATE after education section ─────────────────────────────
    # Cut any cover-letter-style trailing paragraphs after the last
    # "Certification:" line that follows an ACADEMIC QUALIFICATION header.
    # We find the ACADEMIC QUALIFICATION section and then the last cert line after it.
    edu_start = -1
    last_cert_line = -1
    for idx, ln in enumerate(lines):
        lo = ln.strip().lower().rstrip(':')
        if lo in ('academic qualification', 'education & certifications', 'education'):
            edu_start = idx
        if edu_start >= 0 and ln.strip().lower().startswith('certification:'):
            last_cert_line = idx
    if last_cert_line > 0:
        lines = lines[:last_cert_line + 1]

    name_written = False
    contact_done = False
    in_exp       = False
    in_edu       = False
    in_skills    = False
    line_num     = 0

    for raw in lines:
        stripped = raw.strip()
        line_num += 1

        # blank line → small spacer
        if not stripped:
            if name_written:
                p = doc.add_paragraph()
                p.paragraph_format.space_after  = Pt(0)
                p.paragraph_format.space_before = Pt(0)
            continue

        upper = stripped.upper().rstrip(':')

        # ── CANDIDATE NAME (very first real line) ────────────────────────
        if not name_written:
            # Skip literal "HEADER" placeholder the AI sometimes outputs
            if stripped.upper() in ('HEADER', '[HEADER]', '**HEADER**'):
                continue
            name_written = True
            p   = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(stripped.upper())
            rf(run, size=NAME_SIZE, bold=True)
            p.paragraph_format.space_after = Pt(3)
            continue

        # ── CONTACT / HEADLINE (first 6 non-blank lines after name) ──────
        if not contact_done:
            non_blank = [p for p in doc.paragraphs if p.text.strip()]
            if len(non_blank) <= 5:
                is_contact_line = any(k in stripped for k in
                    ['@','Mobile','+65','linkedin','http','|','#0','Road','Street','Avenue'])
                is_headline     = '|' in stripped or stripped.endswith('Excellence') or stripped.endswith('Manager') or stripped.endswith('Owner')
                if is_contact_line or is_headline:
                    p   = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run(stripped)
                    rf(run, bold=True)
                    p.paragraph_format.space_after = Pt(2)
                    continue
                else:
                    contact_done = True
            else:
                contact_done = True

        # ── SECTION HEADERS ──────────────────────────────────────────────
        is_section = (upper in SECTION_HEADERS or
                      (stripped.isupper() and 3 < len(stripped) < 60
                       and not stripped.startswith('-')))
        if is_section:
            in_exp    = 'EXPERIENCE' in upper
            in_edu    = upper in EDUCATION_SECTIONS
            in_skills = 'SKILL' in upper or upper in ('CORE SKILLS',)
            p   = doc.add_paragraph()
            run = p.add_run(stripped.upper().rstrip(':') + ':')
            rf(run, size=HDR_SIZE, bold=True)
            add_border(p)
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after  = Pt(4)
            continue

        # ── BULLET POINTS ────────────────────────────────────────────────
        if stripped.startswith(('- ','* ','– ','• ')):
            content = stripped[2:].strip()
            p   = doc.add_paragraph(style='List Bullet')
            run = p.add_run(content)
            rf(run)
            p.paragraph_format.space_after = Pt(2)
            continue

        # ── SKILL CATEGORY LINES  e.g. "Data visualization tools: Tableau, Power BI"
        skill_match = None
        for cat in SKILL_CATS:
            if stripped.lower().startswith(cat.lower() + ':'):
                skill_match = cat; break
        if skill_match:
            colon = stripped.index(':')
            p   = doc.add_paragraph()
            r1  = p.add_run(stripped[:colon+1])
            rf(r1, bold=True)
            r2  = p.add_run(stripped[colon+1:])
            rf(r2, bold=False)
            p.paragraph_format.space_after = Pt(2)
            continue

        # ── COMPANY + DATE LINE ──────────────────────────────────────────
        is_co_date = bool(re.search(
            r'(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec|\d{4}).{0,40}(\d{4}|Present)',
            stripped, re.I))

        # ── JOB TITLE LINE (bold plain, no bullet) ───────────────────────
        # Must be in experience block, no year, starts with capital or slash,
        # contains a recognisable role keyword, NOT a bullet
        is_job_title = (
            in_exp and
            not is_co_date and
            not stripped.startswith('-') and
            len(stripped) < 100 and
            re.match(r'^[A-Z/\-–]', stripped) and
            any(kw in stripped.lower() for kw in JOB_TITLE_KW) and
            not stripped.upper().rstrip(':') in SECTION_HEADERS
        )

        if is_co_date:
            plain_para(doc, stripped, bold=True, before=5, after=1)
            continue

        if is_job_title:
            plain_para(doc, clean_title(stripped), bold=True, before=0, after=1)
            continue

        # ── EDUCATION BODY TEXT (not bold) ───────────────────────────────
        if in_edu:
            plain_para(doc, stripped, bold=False, before=0, after=3)
            continue

        # ── DEFAULT BODY TEXT ─────────────────────────────────────────────
        plain_para(doc, stripped, bold=False, before=0, after=3)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


@app.route("/api/generate-docs", methods=["POST"])
def generate_docs():
    """Generate tailored resume + cover letter as .docx files using AI + python-docx."""
    import base64

    data = request.json
    role = data.get("role", "").strip()
    company = data.get("company", "").strip()
    jd = data.get("jd", "").strip()
    role_type = data.get("roleType", "").strip()
    matched_keywords = data.get("matchedKeywords", [])

    kw_line = (
        f"PRE-EXTRACTED JD KEYWORDS (use ALL verbatim throughout resume): {', '.join(matched_keywords)}"
        if matched_keywords else
        "Extract and use all keywords from the JD above throughout the resume."
    )

    if not role or not company:
        return jsonify({"error": "role and company are required"}), 400

    try:
        P = get_active_profile()
        ai_role = is_ai_role(jd, role_type)
        framing = build_product_framing(P)

        # Build structured experience text
        exp_text_d = ""
        for exp in P.get('experience', []):
            exp_text_d += f"\n--- {exp.get('role','')} at {exp.get('company','')} ({exp.get('period','')}) ---\n"
            for b in exp.get('bullets', []):
                exp_text_d += f"  - {b}\n"
            for a in exp.get('achievements', []):
                exp_text_d += f"  Achievement: {a}\n"
        skills_text_d = ', '.join(P.get('skills', []))
        edu_text_d = '; '.join([f"{e.get('degree','')} — {e.get('school','')} ({e.get('period','')})" for e in P.get('education', [])])

        # Generate resume text via AI
        resume_prompt = f"""You are an expert ATS resume writer. Write a COMPLETE 2-page resume for {P.get('name','')} targeting: {role_type or role} at {company}.

JOB DESCRIPTION:
{jd[:3000]}

{kw_line}

===== CANDIDATE MASTER DATA =====
Name: {P.get('name','')} | Phone: {P.get('mobile','')} | Email: {P.get('email','')}
LinkedIn: {P.get('linkedin','')} | Address: {P.get('address','')}
Certification: SAFe 6.0 Product Owner/Product Management

HEADLINE OPTIONS (pick most JD-relevant):
- Business Analyst Lead | Product Owner | Sales & Operations Excellence
- Lead Business Analyst / Enterprise Product Owner
- Digital Product Manager | Product Owner (AI & Data Platforms)
- Business Analyst Lead | Product Owner | Product Manager

SUMMARY OPTIONS (blend most relevant):
Option A: Lead Business Analyst with 5+ years supporting senior leadership through business planning, financial operations, and cross-functional project execution. Proven ability to streamline processes, manage executive stakeholder communications, track KPIs, drive operational alignment and cost profitability.
Option B: Lead Business Analyst / Enterprise Product Owner with 5+ years delivering large-scale enterprise platforms across financial services. Proven expertise in owning product roadmaps, translating business needs into scalable technical solutions, and driving platform adoption.
Option C: Digital Product Manager / Product Owner with 5+ years delivering large-scale digital platforms across financial services. Proven expertise in owning product roadmaps, managing Agile delivery, translating user needs into scalable solutions, and driving data-informed product decisions.

SKILLS POOL:
Data visualization tools: Tableau, Power BI
Programming: PSQL, Python basics
Key Modules in Masters: Financial Planning and Analysis, Operations management
Others: Agile, SAFe 6.0, JIRA, Excel, Microsoft Project, Product Vision and road mapping, Business Analysis, Risk Mitigation & Change Management, Budget Forecasting & Variance Analysis, KPI Tracking & Dashboard Reporting, API integrations, Stakeholder Management, AI-enabled Product Integration, SEO, GEO, CRO, MVP Definition, Generative AI, LLM, Prompt Engineering
Certification: Scaled Agile Framework 6.0 Product Owner/Product Management

EXPERIENCE BULLETS POOL — pick 6-8 for KPMG, 3-4 for JPM, 2-3 for Amazon:
KPMG (Feb 2021–Present):
- Defined and drove product vision, roadmap, and delivery strategy for large-scale digital transformation initiatives for financial institutions
- Owned and prioritised product backlog ensuring alignment with business objectives, regulatory requirements and customer experience goals
- Collaborated with cross-regional stakeholders (UX, architecture, finance) to drive alignment across engineering and operations functions
- Drove product scope decisions through impact analysis generating ~5% additional business value
- Led sprint ceremonies (Planning, Reviews, Retros, PI Planning) across multi-squad programme, resulting in 30% increase in team velocity
- Managed 3rd party vendors, conducted go-live planning, led data migrations from legacy systems ensuring 99.9% uptime
- Designed and executed end-to-end test scenarios on Loan IQ applications resulting in 25% reduction in defects
- Supported automation and AI-enabled solution design including API integrations, reducing manual effort by 30 man-days
- Built business cases and executive presentations for new product features securing stakeholder approval
- Drove data-informed product decisions by defining KPIs, analysing performance metrics and identifying improvement opportunities
- Supported executive decision-making through As-Is/To-Be analysis and streamlined documentation of strategic processes
- Trained vendors, managed go-live execution plans, ensuring process continuity across business units
Key Achievements:
- Impact Analysis persuaded clients to approve change requests, generating ~5% additional project profit
- Automated Interest computation solutioning saved 30 man-days of manual work
- Led team through critical sprint-to-SIT transition, maintaining project delivery timeline

J.P. Morgan (Oct 2023–Jan 2024, Asset Management Internship):
- Gathered product requirements from trading/execution teams to build robust investor profiles
- Performed quantitative analysis of 5 stocks, recommended to 2 clients based on risk metrics, achieving 10% portfolio value increase
- Measured portfolio performance via KPIs: Annual Portfolio Return, Variance, Standard Deviation
- Helped clients onboard and build stronger investment portfolios with market-leading solutions

Amazon (Mar 2018–Mar 2019, Business Analyst):
- Built real-time quality monitoring dashboards using Power BI from SQL Server and Excel, reducing quality issues by 20%
- Translated business requirements into functional/non-functional specifications with 95% stakeholder satisfaction
- Analysed and visualised operational data using Tableau and Power BI, increasing operational efficiency 15%

EDUCATION:
Master of Science Engineering Business Management, Coventry University, UK (Jul 2019–Nov 2020)
Bachelor of Engineering Electronics & Communication Engineering, Anna University, India (Jul 2012–Jun 2016)
=================================

OUTPUT FORMAT — FOLLOW EXACTLY, plain text only, no markdown:
Line 1: Amretha Karthikeyan
Line 2: #02-321 153 Gangsa Road, Singapore-670153
Line 3: Mobile: +65-90256503, email: amretha.ammu@gmail.com
Line 4: https://www.linkedin.com/in/amretha-nishanth-534b39101/
Line 5: [chosen headline]
[blank line]
PROFESSIONAL SUMMARY:
[4 sentences blended from options, weaving in JD keywords]
[blank line]
SKILL SET:
Data visualization tools: [tools]
Programming: [languages]
Others: [all relevant skills + JD keywords]
Certification: Scaled Agile Framework 6.0 Product Owner/Product Management
[blank line]
PROFESSIONAL EXPERIENCE:
KPMG, Singapore  Feb 2021 – Present
[Job title — no slashes before or after]
- [6-8 selected bullets]
Key Achievements:
- [3 achievements]
[blank line]
J.P. Morgan  Oct 2023 – Jan 2024
Asset Management Virtual Internship
- [3-4 bullets]
[blank line]
Amazon Inc, India  Mar 2018 – Mar 2019
Business Analyst
- [2-3 bullets]
[blank line]
ACADEMIC QUALIFICATION:
Master of Science Engineering Business Management  Jul 2019 – Nov 2020
Coventry University, UK
Bachelor of Engineering  Jul 2012 – Jun 2016
Electronics & Communication Engineering, Anna University, India
Certification: Scaled Agile Framework 6.0 Product Owner/Product Management

RULES: Plain text only. ALL CAPS section headers. "- " bullets. Job titles on own line with NO slashes. Do NOT write "HEADER". Target 750-850 words. Weave in 12+ exact JD keyword phrases. Do NOT mention the target company name anywhere in the resume. Do NOT add closing paragraphs or cover-letter-style text after the education section."""

        resume_text = call_claude(resume_prompt, max_tokens=8192)
        resume_text = _inject_ai_projects(resume_text)

        # Generate cover letter text via AI
        cover_prompt = f"""Write a professional 300-350 word cover letter for {P['name']} applying to {role_type or role} at {company}.
{framing}

JOB DESCRIPTION:
{jd[:3000]}

{"AI ROLE: Mention AI project experience with URL." if ai_role else ""}

Write in plain text. Be specific about the company and role. Include metrics from candidate experience: ~5% business value, 30 man-days eliminated. Reference SAFe certification."""

        cover_text = call_claude(cover_prompt)

        # Create .docx files
        resume_bytes = _create_docx_from_text(resume_text, f"Resume - {role}")
        cover_bytes = _create_docx_from_text(cover_text, f"Cover Letter - {company}")

        resume_b64 = base64.b64encode(resume_bytes).decode()
        cover_b64 = base64.b64encode(cover_bytes).decode()

        return jsonify({
            "resume_b64": resume_b64,
            "cover_b64": cover_b64,
            "variant": "AI",
            "resume_filename": f"Resume_{company.replace(' ','_')}.docx",
            "cover_filename": f"CoverLetter_{company.replace(' ','_')}.docx",
            "resume_text": resume_text,
            "cover_text": cover_text,
        })
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/import-job", methods=["POST"])
def import_job():
    data = request.json
    url = data.get("url", "").strip()
    
    if not url:
        return jsonify({"error": "No URL provided"}), 400

    # Detect platform
    is_linkedin = "linkedin.com" in url
    is_indeed = "indeed.com" in url or "sg.indeed.com" in url

    # Headers that mimic a real browser
    headers = {
        "User-Agent": "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36",
        "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,image/avif,image/webp,*/*;q=0.8",
        "Accept-Language": "en-US,en;q=0.9",
        "Accept-Encoding": "gzip, deflate, br",
        "Connection": "keep-alive",
        "Upgrade-Insecure-Requests": "1",
        "Cache-Control": "max-age=0",
    }

    result = {
        "platform": "linkedin" if is_linkedin else "indeed" if is_indeed else "other",
        "url": url,
        "title": "",
        "company": "",
        "location": "Singapore",
        "description": "",
        "partial": False,
        "message": ""
    }

    try:
        from bs4 import BeautifulSoup

        if is_linkedin:
            # LinkedIn blocks login-walled pages but public job URLs sometimes work
            # Extract job ID from URL for reference
            import re
            job_id_match = re.search(r'/jobs/view/(\d+)', url)
            job_id = job_id_match.group(1) if job_id_match else ""
            
            # Try to extract company from URL slug
            company_match = re.search(r'linkedin\.com/jobs/view/[^/]+-at-([a-z0-9-]+)-\d+', url)
            if company_match:
                result["company"] = company_match.group(1).replace("-", " ").title()

            try:
                resp = http_requests.get(url, headers=headers, timeout=10)
                soup = BeautifulSoup(resp.text, "lxml")

                # Try various LinkedIn selectors
                title_el = (soup.find("h1", {"class": lambda c: c and "job-title" in c}) or
                           soup.find("h1", {"class": lambda c: c and "topcard__title" in c}) or
                           soup.find("h1"))
                if title_el:
                    result["title"] = title_el.get_text(strip=True)

                company_el = (soup.find("a", {"class": lambda c: c and "topcard__org-name" in c}) or
                             soup.find("span", {"class": lambda c: c and "company-name" in c}))
                if company_el:
                    result["company"] = company_el.get_text(strip=True)

                desc_el = (soup.find("div", {"class": lambda c: c and "description__text" in c}) or
                          soup.find("div", {"class": lambda c: c and "job-description" in c}))
                if desc_el:
                    result["description"] = desc_el.get_text(separator="\n", strip=True)[:3000]

                if not result["title"] and not result["description"]:
                    # LinkedIn returned a login wall
                    result["partial"] = True
                    result["message"] = "LinkedIn requires login to view full details. Company name extracted from URL — please paste the job description manually."
                else:
                    result["message"] = "Job details imported from LinkedIn!"

            except Exception:
                result["partial"] = True
                result["message"] = "LinkedIn blocked the request. Company extracted from URL — please paste the job description manually."

        elif is_indeed:
            resp = http_requests.get(url, headers=headers, timeout=10)
            soup = BeautifulSoup(resp.text, "lxml")

            # Indeed selectors
            title_el = (soup.find("h1", {"class": lambda c: c and "jobTitle" in str(c)}) or
                       soup.find("h1", {"data-testid": "jobsearch-JobInfoHeader-title"}) or
                       soup.find("h1"))
            if title_el:
                result["title"] = title_el.get_text(strip=True).replace("- job post", "").strip()

            company_el = (soup.find("div", {"data-testid": "inlineHeader-companyName"}) or
                         soup.find("span", {"class": lambda c: c and "companyName" in str(c)}) or
                         soup.find("a", {"data-tn-element": "companyName"}))
            if company_el:
                result["company"] = company_el.get_text(strip=True)

            location_el = (soup.find("div", {"data-testid": "job-location"}) or
                          soup.find("div", {"class": lambda c: c and "companyLocation" in str(c)}))
            if location_el:
                result["location"] = location_el.get_text(strip=True)

            desc_el = (soup.find("div", {"id": "jobDescriptionText"}) or
                      soup.find("div", {"class": lambda c: c and "jobsearch-jobDescriptionText" in str(c)}))
            if desc_el:
                result["description"] = desc_el.get_text(separator="\n", strip=True)[:3000]

            if result["title"] or result["company"]:
                result["message"] = "Job details imported from Indeed! ✅"
            else:
                result["partial"] = True
                result["message"] = "Could not extract details automatically. Please fill in manually."

        else:
            # Generic scrape attempt
            resp = http_requests.get(url, headers=headers, timeout=10)
            soup = BeautifulSoup(resp.text, "lxml")
            title_el = soup.find("h1")
            if title_el:
                result["title"] = title_el.get_text(strip=True)
            result["partial"] = True
            result["message"] = "Basic details extracted — please verify and fill in any missing fields."

    except Exception as e:
        result["partial"] = True
        result["message"] = f"Could not fetch URL automatically. Please fill in details manually. ({str(e)[:80]})"

    return jsonify(result)


@app.route("/api/rank-jobs", methods=["POST"])
def rank_jobs():
    """
    JD keyword extraction + profile match scoring. No AI API — instant, never fails.
    Also returns matched_keywords per job so resume generation can reuse them.

    Scoring (max 10):
      Keyword match score   0–4.5 pts  (weighted by category)
      Role type match       0–2.0 pts
      Company bonus/penalty ±2.0 pts
      Product co signals    0–1.0 pt
      WLB signals           0–0.5 pt
      Singapore location    0–0.5 pt
      Visa sponsorship      0–0.5 pt  (+override to 0 if blocked)
    """
    import re as _re

    data = request.json or {}
    jobs = data.get("jobs", [])
    if not jobs:
        return jsonify({"error": "No jobs provided"}), 400

    P = get_active_profile()

    # ── Candidate profile text (for matching) ────────────────────────────
    profile_text = " ".join([
        P.get("summary", ""),
        P.get("headline", ""),
        " ".join(P.get("skills", [])),
        " ".join(
            b for exp in P.get("experience", [])
            for b in exp.get("bullets", []) + exp.get("achievements", [])
        ),
        # hardcode known skills not always in profile
        "agile safe scrum kanban jira confluence sql python tableau power bi api "
        "roadmap backlog user stories sprint stakeholder kpi dashboard fintech "
        "banking payments digital transformation product roadmap mvp uat change management "
        "business analysis product owner product manager ba pm po saas b2b b2c "
        "generative ai llm prompt engineering flask supabase render",
    ]).lower()

    # ── Keyword categories with weights ──────────────────────────────────
    # (keyword, weight, category_label)
    KEYWORD_DEFS = [
        # Role keywords — weight 0.6
        ("product owner",           0.6, "role"),
        ("product manager",         0.6, "role"),
        ("product management",      0.5, "role"),
        ("business analyst",        0.6, "role"),
        ("product lead",            0.5, "role"),
        ("product operations",      0.5, "role"),
        ("delivery manager",        0.4, "role"),
        ("scrum master",            0.4, "role"),
        ("agile coach",             0.4, "role"),
        # Methodology keywords — weight 0.35
        ("agile",                   0.35, "method"),
        ("safe",                    0.35, "method"),
        ("scrum",                   0.35, "method"),
        ("kanban",                  0.35, "method"),
        ("sprint",                  0.3,  "method"),
        ("pi planning",             0.35, "method"),
        # Tools — weight 0.3
        ("jira",                    0.3,  "tool"),
        ("confluence",              0.3,  "tool"),
        ("sql",                     0.3,  "tool"),
        ("python",                  0.3,  "tool"),
        ("tableau",                 0.3,  "tool"),
        ("power bi",                0.3,  "tool"),
        ("figma",                   0.25, "tool"),
        ("miro",                    0.25, "tool"),
        ("excel",                   0.2,  "tool"),
        ("notion",                  0.2,  "tool"),
        ("amplitude",               0.3,  "tool"),
        ("mixpanel",                0.3,  "tool"),
        ("looker",                  0.3,  "tool"),
        # Product skills — weight 0.35
        ("product roadmap",         0.35, "skill"),
        ("roadmap",                 0.3,  "skill"),
        ("backlog",                 0.35, "skill"),
        ("user stories",            0.35, "skill"),
        ("stakeholder management",  0.35, "skill"),
        ("stakeholder",             0.25, "skill"),
        ("kpi",                     0.3,  "skill"),
        ("dashboard",               0.25, "skill"),
        ("data analysis",           0.3,  "skill"),
        ("data-driven",             0.3,  "skill"),
        ("go-to-market",            0.35, "skill"),
        ("gtm",                     0.3,  "skill"),
        ("mvp",                     0.3,  "skill"),
        ("uat",                     0.3,  "skill"),
        ("change management",       0.3,  "skill"),
        ("product vision",          0.35, "skill"),
        ("discovery",               0.3,  "skill"),
        ("a/b testing",             0.3,  "skill"),
        ("experimentation",         0.3,  "skill"),
        ("api",                     0.25, "skill"),
        ("api integration",         0.3,  "skill"),
        ("requirements",            0.25, "skill"),
        ("business case",           0.25, "skill"),
        ("seo",                     0.3,  "skill"),
        ("cro",                     0.3,  "skill"),
        ("ux",                      0.25, "skill"),
        ("user research",           0.3,  "skill"),
        ("customer journey",        0.3,  "skill"),
        # Domain — weight 0.3
        ("fintech",                 0.35, "domain"),
        ("banking",                 0.3,  "domain"),
        ("payments",                0.3,  "domain"),
        ("digital transformation",  0.3,  "domain"),
        ("saas",                    0.3,  "domain"),
        ("b2b",                     0.25, "domain"),
        ("b2c",                     0.25, "domain"),
        ("platform",                0.2,  "domain"),
        ("financial services",      0.3,  "domain"),
        ("e-commerce",              0.25, "domain"),
        ("marketplace",             0.25, "domain"),
        # AI/ML — weight 0.4
        ("generative ai",           0.4,  "ai"),
        ("llm",                     0.4,  "ai"),
        ("machine learning",        0.35, "ai"),
        ("ai",                      0.3,  "ai"),
        ("prompt engineering",      0.4,  "ai"),
        ("nlp",                     0.35, "ai"),
    ]

    BONUS_COMPANIES = [
        "grab", "sea", "shopee", "gojek", "airwallex", "stripe", "revolut",
        "wise", "transferwise", "propertyguru", "carousell", "govtech",
        "dbs", "ocbc", "uob", "singlife", "nium", "rapyd", "aspire",
        "lazada", "bytedance", "tiktok", "foodpanda", "delivery hero",
        "google", "meta", "netflix", "amazon", "apple", "microsoft",
    ]
    PENALTY_COMPANIES = [
        "kpmg", "deloitte", "pwc", "ey ", "ernst", "accenture", "mckinsey",
        "bcg", "bain", "ibm", "wipro", "infosys", "tcs", "cognizant",
        "capgemini", "ncs ", "dxc", "fujitsu",
    ]
    VISA_BLOCK_PHRASES = [
        "no sponsorship", "citizen or pr", "citizen/pr", "pr only", "citizens only",
        "must be a citizen", "must be singapore", "singaporean only",
        "ep not provided", "no ep", "no work pass", "own ep",
        "singapore citizens and pr", "singapore citizen or pr",
        "must hold", "only singaporean",
    ]
    VISA_OK_PHRASES = [
        "visa sponsorship", "ep sponsorship", "work pass", "s pass", "ep provided",
        "sponsorship available", "open to ep", "employment pass provided",
        "relocation support", "open to all nationalities",
    ]
    WLB_SIGNALS = [
        "work life balance", "work-life", "flexible", "hybrid", "remote",
        "well-being", "wellness", "benefits", "learning & development",
        "flat structure", "transparent", "autonomy", "ownership culture",
    ]
    PRODUCT_CO_SIGNALS = [
        "product-led", "product company", "in-house", "saas", "platform team",
        "b2b", "b2c", "consumer product", "marketplace", "fintech",
        "proptech", "healthtech", "edtech", "our product", "we build",
        "product organisation", "product org",
    ]
    SG_SIGNALS = [
        "singapore", " sg ", "sgd", "raffles", "tanjong pagar",
        "one-north", "mapletree", "mbfc", "orchard", "marina bay",
    ]

    rankings = []

    for j in jobs:
        jd_raw  = (j.get("jd") or "")
        role    = (j.get("role") or "").lower()
        company = (j.get("company") or "").lower()
        jd      = jd_raw.lower()
        combined = jd + " " + role + " " + company

        score   = 0.0
        reasons = []
        matched_keywords = []   # reused by resume generation

        # ── 1. Visa hard override ─────────────────────────────────────────
        visa_blocked = any(p in combined for p in VISA_BLOCK_PHRASES)
        visa_ok      = any(p in combined for p in VISA_OK_PHRASES)
        if visa_blocked and not visa_ok:
            rankings.append({
                "id":               j.get("id"),
                "score":            0,
                "label":            "❌ Weak Fit",
                "priority":         "Skip",
                "reason":           "No visa sponsorship — requires Singapore Citizen/PR only.",
                "matched_keywords": [],
            })
            continue

        # ── 2. Extract JD keywords & match against profile ───────────────
        # For each keyword in our dictionary: check if it appears in JD AND in profile
        jd_present      = []  # keywords found in JD
        profile_matched = []  # keywords found in both JD + profile
        kw_score        = 0.0

        for kw, weight, cat in KEYWORD_DEFS:
            if kw in jd or kw in combined:
                jd_present.append((kw, weight, cat))
                if kw in profile_text:
                    profile_matched.append((kw, weight, cat))
                    kw_score += weight

        kw_score = min(4.5, kw_score)
        score += kw_score

        # Top matched keywords for display + resume reuse
        matched_keywords = [kw for kw, _, _ in sorted(profile_matched, key=lambda x: -x[1])[:15]]
        jd_only_keywords = [kw for kw, _, _ in jd_present if kw not in matched_keywords][:8]

        if profile_matched:
            top5 = ", ".join(matched_keywords[:5])
            reasons.append(f"Profile matches {len(profile_matched)} JD keywords: {top5}")

        # ── 3. Role type match (0–2 pts) ─────────────────────────────────
        role_kws = [kw for kw, _, cat in profile_matched if cat == "role"]
        jd_role_kws = [kw for kw, _, cat in jd_present if cat == "role"]
        role_score = min(2.0, len(jd_role_kws) * 0.7)
        score += role_score
        if jd_role_kws:
            reasons.append(f"Role: {jd_role_kws[0]}")

        # ── 4. Company bonus / penalty ────────────────────────────────────
        co_bonus   = any(b in company for b in BONUS_COMPANIES)
        co_penalty = any(p in company for p in PENALTY_COMPANIES)
        if co_bonus:
            score += 2
            reasons.append(f"Target company ✓")
        elif co_penalty:
            score = min(score, 4.0)
            reasons.append("Consulting firm — capped")

        # ── 5. Product company signals (+1) ──────────────────────────────
        if any(s in combined for s in PRODUCT_CO_SIGNALS):
            score += 1
            reasons.append("Product/in-house company")

        # ── 6. WLB signals (+0.5) ────────────────────────────────────────
        if any(s in combined for s in WLB_SIGNALS):
            score += 0.5
            reasons.append("Good WLB signals")

        # ── 7. Singapore location (+0.5) ─────────────────────────────────
        if any(s in combined for s in SG_SIGNALS):
            score += 0.5

        # ── 8. Visa sponsorship offered (+0.5) ───────────────────────────
        if visa_ok:
            score += 0.5
            reasons.append("Visa/EP sponsorship available ✓")

        score = round(min(10.0, max(1.0, score)), 1)

        # ── Label + priority ─────────────────────────────────────────────
        if score >= 8:
            label, priority = "🔥 Strong Match", "Apply Today"
        elif score >= 6.5:
            label, priority = "✅ Good Fit",     "Apply This Week"
        elif score >= 4.5:
            label, priority = "🟡 Possible",     "Lower Priority"
        else:
            label, priority = "❌ Weak Fit",     "Skip"

        reason_str = ". ".join(reasons[:3]) if reasons else "Based on role and keyword analysis."

        rankings.append({
            "id":               j.get("id"),
            "score":            score,
            "label":            label,
            "priority":         priority,
            "reason":           reason_str,
            "matched_keywords": matched_keywords,      # for resume generation
            "jd_only_keywords": jd_only_keywords,      # keywords in JD not yet in profile
        })

    print(f"[rank_jobs] Scored {len(rankings)} jobs (keyword match, no AI)")
    return jsonify({"rankings": rankings})


@app.route("/api/fetch-jd", methods=["POST"])
def fetch_jd():
    """Fetch job description from a LinkedIn or Indeed URL via HTTP scraping.
    For LinkedIn URLs, tries Voyager API with stored li_at cookie first."""
    import requests as req
    import uuid as _uuid
    from bs4 import BeautifulSoup

    data = request.json
    url = (data.get("url") or "").strip()
    if not url:
        return jsonify({"error": "No URL provided"}), 400

    try:
        jd = ""
        title = ""
        company = ""

        # ── LinkedIn: try Voyager API with stored cookie first ──
        if "linkedin.com" in url:
            m = __import__("re").search(r"/jobs/view/(\d+)", url)
            if m:
                job_id = m.group(1)
                li_at = _get_li_at_cookie()
                if li_at:
                    li_at = li_at.strip().strip('"').strip("'")
                    csrf_token = f"ajax:{_uuid.uuid4()}"
                    v_headers = {
                        "User-Agent": "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) "
                                      "AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36",
                        "Accept": "application/vnd.linkedin.normalized+json+2.1",
                        "x-restli-protocol-version": "2.0.0",
                        "csrf-token": csrf_token,
                    }
                    v_cookies = {"li_at": li_at, "JSESSIONID": f'"{csrf_token}"'}
                    v_url = (
                        f"https://www.linkedin.com/voyager/api/jobs/jobPostings/{job_id}"
                        f"?decorationId=com.linkedin.voyager.deco.jobs.web.shared.WebFullJobPosting-65"
                    )
                    try:
                        vr = http_requests.get(v_url, headers=v_headers, cookies=v_cookies, timeout=15)
                        if vr.status_code == 200:
                            vd = vr.json()
                            desc = vd.get("description") or vd.get("descriptionText") or {}
                            if isinstance(desc, dict):
                                jd = desc.get("text", "")[:5000]
                            elif isinstance(desc, str):
                                jd = desc[:5000]
                            # Also get title/company from Voyager response
                            title = vd.get("title") or ""
                            comp_detail = vd.get("companyDetails") or {}
                            if isinstance(comp_detail, dict):
                                comp_res = comp_detail.get("com.linkedin.voyager.deco.jobs.web.shared.WebJobPostingCompany") or comp_detail
                                company = comp_res.get("companyResolutionResult", {}).get("name") or comp_res.get("company", {}).get("name") or ""
                            if jd:
                                return jsonify({"jd": jd, "title": title, "company": company})
                    except Exception as ve:
                        print(f"[fetch-jd] Voyager API failed for {job_id}: {ve}")

        # ── Fallback: public HTML scraping ──
        headers = {
            "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36",
            "Accept": "text/html,application/xhtml+xml",
            "Accept-Language": "en-US,en;q=0.9",
        }
        resp = req.get(url, headers=headers, timeout=15)
        soup = BeautifulSoup(resp.text, "html.parser")

        if "linkedin.com" in url:
            # LinkedIn job description selectors
            for sel in [
                ".description__text",
                ".show-more-less-html__markup",
                "[class*='description']",
                "section.description",
            ]:
                el = soup.select_one(sel)
                if el and len(el.get_text(strip=True)) > 100:
                    jd = el.get_text(separator="\n", strip=True)[:5000]
                    break

            title_el = soup.select_one("h1.top-card-layout__title, h1[class*='title']")
            if title_el:
                title = title_el.get_text(strip=True)

            company_el = soup.select_one("a.topcard__org-name-link, [class*='company-name']")
            if company_el:
                company = company_el.get_text(strip=True)

        elif "indeed.com" in url:
            for sel in ["#jobDescriptionText", ".jobsearch-jobDescriptionText", "[class*='description']"]:
                el = soup.select_one(sel)
                if el and len(el.get_text(strip=True)) > 100:
                    jd = el.get_text(separator="\n", strip=True)[:5000]
                    break

        elif "mycareersfuture.gov.sg" in url:
            for sel in ["[class*='job-description']", "[class*='description']", "article"]:
                el = soup.select_one(sel)
                if el and len(el.get_text(strip=True)) > 100:
                    jd = el.get_text(separator="\n", strip=True)[:5000]
                    break
        else:
            # Generic fallback
            for tag in soup.find_all(["article", "section", "div"], limit=20):
                text = tag.get_text(strip=True)
                if len(text) > 500 and any(kw in text.lower() for kw in ["responsibilities", "requirements", "qualifications", "experience"]):
                    jd = text[:5000]
                    break

        if not jd:
            return jsonify({"jd": "", "error": "Could not extract JD — LinkedIn may require login"}), 200

        return jsonify({"jd": jd, "title": title, "company": company})

    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/bookmarklet-add", methods=["POST", "OPTIONS"])
def bookmarklet_add():
    # Handle CORS preflight - bookmarklet calls come from linkedin.com/indeed.com
    if request.method == "OPTIONS":
        response = app.make_default_options_response()
        response.headers["Access-Control-Allow-Origin"] = "*"
        response.headers["Access-Control-Allow-Headers"] = "Content-Type"
        response.headers["Access-Control-Allow-Methods"] = "POST, OPTIONS"
        return response

    data = request.json
    role = data.get("role", "").strip()
    company = data.get("company", "").strip()

    if not role or not company:
        resp = jsonify({"success": False, "error": "Missing job title or company"})
        resp.headers["Access-Control-Allow-Origin"] = "*"
        return resp, 400

    # Store in a simple JSON file on the server
    import json, os
    jobs_file = os.path.join(BASE_DIR, "bookmarked_jobs.json")
    
    try:
        if os.path.exists(jobs_file):
            with open(jobs_file, "r") as f:
                saved = json.load(f)
        else:
            saved = []

        new_job = {
            "id": int(__import__("time").time() * 1000),
            "role": role,
            "company": company,
            "jd": data.get("jd", ""),
            "location": data.get("location", "Singapore"),
            "url": data.get("url", ""),
            "status": "wishlist",
            "date": __import__("datetime").date.today().strftime("%d/%m/%Y"),
            "notes": "",
            "salary": "",
            "isDemo": False,
            "fromBookmarklet": True
        }

        saved.append(new_job)
        with open(jobs_file, "w") as f:
            json.dump(saved, f)

        resp = jsonify({"success": True, "job": new_job})
        resp.headers["Access-Control-Allow-Origin"] = "*"
        return resp

    except Exception as e:
        resp = jsonify({"success": False, "error": str(e)})
        resp.headers["Access-Control-Allow-Origin"] = "*"
        return resp, 500


@app.route("/api/pending-count", methods=["GET"])
def pending_count():
    """Returns how many jobs are queued — does NOT clear the file"""
    import json, os
    jobs_file = os.path.join(BASE_DIR, "bookmarked_jobs.json")
    if not os.path.exists(jobs_file):
        return jsonify({"count": 0})
    try:
        saved = json.load(open(jobs_file))
        return jsonify({"count": len(saved)})
    except:
        return jsonify({"count": 0})


@app.route("/api/bookmarklet-jobs", methods=["GET"])
def bookmarklet_jobs():
    """Frontend calls this to pull queued jobs — clears file after sending"""
    import json, os
    jobs_file = os.path.join(BASE_DIR, "bookmarked_jobs.json")
    if not os.path.exists(jobs_file):
        return jsonify({"jobs": []})
    try:
        with open(jobs_file, "r") as f:
            saved = json.load(f)
        with open(jobs_file, "w") as f:
            json.dump([], f)
        return jsonify({"jobs": saved})
    except Exception as e:
        return jsonify({"jobs": [], "error": str(e)})


@app.route("/api/bookmarklet-bulk", methods=["POST", "OPTIONS"])
def bookmarklet_bulk():
    if request.method == "OPTIONS":
        response = app.make_default_options_response()
        response.headers["Access-Control-Allow-Origin"] = "*"
        response.headers["Access-Control-Allow-Headers"] = "Content-Type"
        response.headers["Access-Control-Allow-Methods"] = "POST, OPTIONS"
        return response

    data = request.json
    incoming_jobs = data.get("jobs", [])

    if not incoming_jobs:
        resp = jsonify({"success": False, "error": "No jobs provided"})
        resp.headers["Access-Control-Allow-Origin"] = "*"
        return resp, 400

    import json, os, time
    from datetime import date

    jobs_file = os.path.join(BASE_DIR, "bookmarked_jobs.json")

    try:
        if os.path.exists(jobs_file):
            with open(jobs_file, "r") as f:
                saved = json.load(f)
        else:
            saved = []

        # Avoid duplicates by URL or title+company combo
        seen_urls = set()
        seen_tc   = set()
        for j in saved:
            u = j.get("url","").split("?")[0]
            tc = (j.get("role","").strip().lower() + "|" + j.get("company","").strip().lower())
            if u: seen_urls.add(u)
            seen_tc.add(tc)

        added = 0
        for job in incoming_jobs:
            u  = job.get("url","").split("?")[0]
            tc = (job.get("role","").strip().lower() + "|" + job.get("company","").strip().lower())
            if (u and u in seen_urls) or tc in seen_tc:
                continue
            new_job = {
                "id": int(time.time() * 1000) + added,
                "role": job.get("role", "").strip(),
                "company": job.get("company", "").strip(),
                "jd": job.get("jd", ""),
                "location": job.get("location", "Singapore"),
                "url": job.get("url", ""),
                "status": "wishlist",
                "date": date.today().strftime("%d/%m/%Y"),
                "notes": "",
                "salary": "",
                "isDemo": False,
                "fromBookmarklet": True
            }
            saved.append(new_job)
            if u: seen_urls.add(u)
            seen_tc.add(tc)
            added += 1

        with open(jobs_file, "w") as f:
            json.dump(saved, f)

        resp = jsonify({"success": True, "count": added, "total": len(incoming_jobs)})
        resp.headers["Access-Control-Allow-Origin"] = "*"
        return resp

    except Exception as e:
        resp = jsonify({"success": False, "error": str(e)})
        resp.headers["Access-Control-Allow-Origin"] = "*"
        return resp, 500


@app.route("/ping")
def ping():
    return "ok", 200


@app.route("/capture")
def capture():
    """Bookmarklet redirects here with job details as URL params"""
    title = request.args.get("title", "").strip()
    company = request.args.get("company", "").strip()
    location = request.args.get("location", "Singapore").strip()
    url = request.args.get("url", "").strip()
    jd = request.args.get("jd", "").strip()

    if not title:
        title = "Unknown Role"
    if not company:
        company = "Unknown Company"

    import json, os, time
    from datetime import date

    jobs_file = os.path.join(BASE_DIR, "bookmarked_jobs.json")
    try:
        existing = json.load(open(jobs_file)) if os.path.exists(jobs_file) else []
    except:
        existing = []

    # Dedup check
    clean_url = url.split("?")[0]
    already = any(j.get("url","").split("?")[0] == clean_url or
                  (j.get("role","") == title and j.get("company","") == company)
                  for j in existing)

    if not already:
        existing.append({
            "id": int(time.time() * 1000),
            "role": title,
            "company": company,
            "location": location,
            "url": url,
            "jd": jd,
            "status": "wishlist",
            "date": date.today().strftime("%d/%m/%Y"),
            "notes": "",
            "salary": "",
            "isDemo": False,
            "fromBookmarklet": True
        })
        with open(jobs_file, "w") as f:
            json.dump(existing, f)
        msg = f"✅ <strong>{title}</strong> at <strong>{company}</strong> added to your Job Tracker!"
        color = "#15803d"
    else:
        msg = f"⚠️ <strong>{title}</strong> at <strong>{company}</strong> is already in your tracker."
        color = "#c2410c"

    app_url = request.host_url.rstrip('/')
    return f"""<!DOCTYPE html>
<html><head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>Job Saved!</title>
<style>
  body {{ font-family: -apple-system, sans-serif; background: #f8fafc; display: flex; align-items: center; justify-content: center; min-height: 100vh; margin: 0; }}
  .card {{ background: white; border-radius: 16px; padding: 40px; max-width: 420px; width: 90%; box-shadow: 0 4px 24px rgba(0,0,0,0.1); text-align: center; }}
  h2 {{ color: {color}; margin-bottom: 8px; font-size: 22px; }}
  p {{ color: #64748b; margin-bottom: 24px; font-size: 15px; line-height: 1.5; }}
  .btn {{ display: inline-block; padding: 12px 24px; border-radius: 8px; text-decoration: none; font-weight: 600; font-size: 15px; cursor: pointer; border: none; margin: 6px; }}
  .btn-primary {{ background: #6366f1; color: white; }}
  .btn-ghost {{ background: #f1f5f9; color: #475569; }}
</style>
</head><body>
<div class="card">
  <div style="font-size:48px;margin-bottom:16px;">{'🎯' if not already else '📌'}</div>
  <h2>{'Job Saved!' if not already else 'Already Saved'}</h2>
  <p>{msg}</p>
  <a href="{app_url}" class="btn btn-primary">Open Job Tracker</a>
  <button onclick="history.back()" class="btn btn-ghost">← Back to LinkedIn</button>
</div>
<script>
  // Auto-close and go back to LinkedIn after 3 seconds if opened in same tab
  setTimeout(function() {{ history.back(); }}, 3000);
</script>
</body></html>"""


@app.route("/capture-bulk", methods=["POST"])
def capture_bulk():
    import json, os, time
    from datetime import date

    raw = request.form.get("jobs", "[]")
    try:
        incoming = json.loads(raw)
    except:
        return "Invalid data", 400

    jobs_file = os.path.join(BASE_DIR, "bookmarked_jobs.json")
    try:
        existing = json.load(open(jobs_file)) if os.path.exists(jobs_file) else []
    except:
        existing = []

    # Build lookup sets for both URL and title+company
    seen_urls = set()
    seen_tc   = set()
    for j in existing:
        u = j.get("url","").split("?")[0]
        tc = (j.get("role","").strip().lower() + "|" + j.get("company","").strip().lower())
        if u: seen_urls.add(u)
        seen_tc.add(tc)

    added = 0
    for job in incoming:
        u  = job.get("url","").split("?")[0]
        tc = (job.get("role","").strip().lower() + "|" + job.get("company","").strip().lower())
        if (u and u in seen_urls) or tc in seen_tc:
            continue
        existing.append({
            "id": int(time.time() * 1000) + added,
            "role": job.get("role","").strip(),
            "company": job.get("company","").strip(),
            "location": job.get("location","Singapore"),
            "url": job.get("url",""),
            "jd": "",
            "status": "wishlist",
            "roleType": job.get("roleType","Business Analyst"),
            "priority": job.get("priority","Medium"),
            "source": "LinkedIn",
            "dateApplied": date.today().isoformat(),
            "notes": "", "salary": "", "isDemo": False,
            "fromBookmarklet": True, "checklist": {}
        })
        if u: seen_urls.add(u)
        seen_tc.add(tc)
        added += 1

    with open(jobs_file, "w") as f:
        json.dump(existing, f)

    # Redirect back to the tracker — user lands there and clicks "Import Pending Jobs"
    return redirect(f"/?imported={added}")


# ═══════════════════════════════════════════════════════════════
# JOB DISCOVERY — Scrape LinkedIn, Workable, MyCareersFuture + more
# ═══════════════════════════════════════════════════════════════

import urllib.parse
import re as _re
import datetime as _dt
from concurrent.futures import ThreadPoolExecutor, as_completed


def _scrape_mycareersfuture(keywords, location, max_days):
    """Scrape MyCareersFuture.gov.sg — Singapore Government job portal (free API, no auth)."""
    jobs = []
    try:
        query = urllib.parse.quote_plus(keywords)
        # MCF API supports pagination (limit up to 100), sorted by newest
        api_url = f"https://api.mycareersfuture.gov.sg/v2/jobs?search={query}&limit=50&page=0&sortBy=new_posting_date"

        resp = http_requests.get(api_url, headers={
            "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36",
            "Accept": "application/json",
        }, timeout=20)

        if resp.status_code == 200:
            data = resp.json()
            for item in data.get("results", []):
                title = item.get("title", "")
                company = (item.get("postedCompany") or {}).get("name", "")
                url = item.get("metadata", {}).get("jobDetailsUrl", "")
                if not url and item.get("uuid"):
                    url = f"https://www.mycareersfuture.gov.sg/job/{item['uuid']}"

                # Salary info
                sal = item.get("salary", {})
                if isinstance(sal, dict):
                    sal_min = sal.get("minimum", "")
                    sal_max = sal.get("maximum", "")
                    sal_type = (sal.get("type") or {}).get("salaryType", "") if isinstance(sal.get("type"), dict) else ""
                    salary_str = f"${sal_min}-${sal_max} {sal_type}".strip() if sal_min else ""
                else:
                    salary_str = ""

                # Description (HTML) — strip tags for plain text
                desc_html = item.get("description", "") or ""
                from bs4 import BeautifulSoup
                desc_text = BeautifulSoup(desc_html, "html.parser").get_text(separator=" ", strip=True)[:4000]

                # Posted date → days ago
                posted = item.get("metadata", {}).get("newPostingDate", "")
                days_ago = None
                if posted:
                    try:
                        pd = _dt.datetime.strptime(posted[:10], "%Y-%m-%d")
                        days_ago = (_dt.datetime.now() - pd).days
                    except Exception:
                        pass

                if title:
                    jobs.append({
                        "role": title,
                        "company": company,
                        "url": url,
                        "location": "Singapore",
                        "jd": desc_text,
                        "salary": salary_str,
                        "platform": "MyCareersFuture",
                        "postedDaysAgo": days_ago,
                    })

            # Filter by max_days if we have date info
            if max_days and jobs:
                jobs = [j for j in jobs if j.get("postedDaysAgo") is None or j["postedDaysAgo"] <= max_days]

    except Exception as e:
        return jobs, str(e)

    return jobs, None


def _scrape_linkedin_guest(keywords, location, max_days):
    """Scrape LinkedIn using the public guest Jobs API (no auth needed, paginated)."""
    jobs = []
    try:
        from bs4 import BeautifulSoup
        query = urllib.parse.quote_plus(keywords)
        loc = urllib.parse.quote_plus(location)
        time_filter = "r86400" if max_days <= 1 else "r604800" if max_days <= 7 else "r2592000"

        headers = {
            "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/124.0.0.0 Safari/537.36",
            "Accept": "text/html",
        }

        # LinkedIn guest API returns 10 per page — fetch 3 pages = 30 jobs
        for start in [0, 10, 20]:
            try:
                url = (f"https://www.linkedin.com/jobs-guest/jobs/api/seeMoreJobPostings/search"
                       f"?keywords={query}&location={loc}&f_TPR={time_filter}&start={start}")
                resp = http_requests.get(url, headers=headers, timeout=15)
                if resp.status_code != 200:
                    continue
                soup = BeautifulSoup(resp.text, "html.parser")
                cards = soup.select("div.base-card, li.result-card, div.job-search-card")
                for card in cards:
                    try:
                        title_el = card.select_one("h3.base-search-card__title, h3[class*='title']")
                        title = title_el.get_text(strip=True) if title_el else ""
                        company_el = card.select_one("h4.base-search-card__subtitle, a[class*='company']")
                        company = company_el.get_text(strip=True) if company_el else ""
                        link_el = card.select_one("a.base-card__full-link, a[class*='job-card']")
                        href = link_el.get("href", "") if link_el else ""
                        loc_el = card.select_one("span.job-search-card__location")
                        job_loc = loc_el.get_text(strip=True) if loc_el else location
                        time_el = card.select_one("time")
                        days_ago = None
                        if time_el and time_el.get("datetime"):
                            try:
                                pd = _dt.datetime.strptime(time_el["datetime"][:10], "%Y-%m-%d")
                                days_ago = (_dt.datetime.now() - pd).days
                            except Exception:
                                pass

                        if title:
                            jobs.append({
                                "role": title,
                                "company": company,
                                "url": href.split("?")[0] if href else "",
                                "location": job_loc,
                                "jd": "",
                                "platform": "LinkedIn+",
                                "postedDaysAgo": days_ago,
                            })
                    except Exception:
                        continue
            except Exception:
                continue

    except Exception as e:
        return jobs, str(e)

    return jobs, None


def _scrape_workable(keywords, location, max_days):
    """Search Workable job board using their search API."""
    jobs = []
    try:
        # Workable has a search API
        api_url = "https://jobs.workable.com/api/v1/jobs"
        params = {
            "query": keywords,
            "location": location,
            "limit": 30,
        }
        headers = {
            "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36",
            "Accept": "application/json",
        }

        try:
            resp = http_requests.get(api_url, params=params, headers=headers, timeout=15)
            if resp.status_code == 200:
                data = resp.json()
                job_items = data if isinstance(data, list) else data.get("results", data.get("jobs", []))
                for item in (job_items or [])[:30]:
                    if isinstance(item, dict):
                        title = item.get("title", item.get("name", ""))
                        company = item.get("company", item.get("organization", {}) if isinstance(item.get("organization"), dict) else {})
                        if isinstance(company, dict):
                            company = company.get("name", "") or company.get("title", "")
                        job_url = item.get("url", item.get("application_url", ""))
                        job_loc = item.get("location", location)
                        if isinstance(job_loc, dict):
                            job_loc = job_loc.get("city", location)
                        jd = (item.get("description", "") or "")[:2000]

                        if title:
                            jobs.append({
                                "role": title,
                                "company": company if isinstance(company, str) else "",
                                "url": job_url,
                                "location": job_loc if isinstance(job_loc, str) else location,
                                "jd": jd,
                                "platform": "Workable",
                                "postedDaysAgo": None,
                            })
        except Exception:
            pass

        # Fallback: HTML scraping
        if not jobs:
            from bs4 import BeautifulSoup
            query = urllib.parse.quote_plus(keywords)
            loc = urllib.parse.quote_plus(location)
            url = f"https://jobs.workable.com/?query={query}&location={loc}"
            try:
                resp = http_requests.get(url, headers={
                    "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/124.0.0.0 Safari/537.36",
                    "Accept": "text/html",
                }, timeout=15)
                soup = BeautifulSoup(resp.text, "html.parser")

                # Try JSON-LD
                for script in soup.find_all("script", type="application/ld+json"):
                    try:
                        ld_data = json.loads(script.string or "")
                        if isinstance(ld_data, list):
                            for jp in ld_data:
                                if jp.get("@type") == "JobPosting" and jp.get("title"):
                                    jobs.append({
                                        "role": jp["title"],
                                        "company": (jp.get("hiringOrganization", {}) or {}).get("name", ""),
                                        "url": jp.get("url", ""),
                                        "location": location,
                                        "jd": (jp.get("description", "") or "")[:2000],
                                        "platform": "Workable",
                                        "postedDaysAgo": None,
                                    })
                    except Exception:
                        continue

                # Try links
                if not jobs:
                    all_links = soup.select("a[href*='/j/'], a[href*='/view/']")
                    seen = set()
                    for link in all_links[:30]:
                        href = link.get("href", "")
                        if href and href not in seen:
                            seen.add(href)
                            title_text = link.get_text(strip=True)
                            if 3 < len(title_text) < 200:
                                full_url = href if href.startswith("http") else f"https://jobs.workable.com{href}"
                                jobs.append({
                                    "role": title_text,
                                    "company": "",
                                    "url": full_url,
                                    "location": location,
                                    "jd": "",
                                    "platform": "Workable",
                                    "postedDaysAgo": None,
                                })
            except Exception:
                pass

    except Exception as e:
        return jobs, str(e)

    return jobs, None


def _scrape_linkedin_public(keywords, location, max_days):
    """Search LinkedIn public job listings (no login required)."""
    jobs = []
    try:
        query = urllib.parse.quote_plus(keywords)
        loc = urllib.parse.quote_plus(location)
        # LinkedIn public jobs search — f_TPR=r2592000 = last 30 days
        time_filter = "r86400" if max_days <= 1 else "r604800" if max_days <= 7 else "r2592000"
        url = f"https://www.linkedin.com/jobs/search/?keywords={query}&location={loc}&f_TPR={time_filter}&position=1&pageNum=0"

        headers = {
            "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/124.0.0.0 Safari/537.36",
            "Accept": "text/html,application/xhtml+xml",
            "Accept-Language": "en-US,en;q=0.9",
        }

        from bs4 import BeautifulSoup
        resp = http_requests.get(url, headers=headers, timeout=20, allow_redirects=True)
        if resp.status_code == 200:
            soup = BeautifulSoup(resp.text, "html.parser")
            cards = soup.select("div.base-card, li.result-card, div.job-search-card")
            for card in cards[:30]:
                try:
                    title_el = card.select_one("h3.base-search-card__title, h3[class*='title']")
                    title = title_el.get_text(strip=True) if title_el else ""
                    company_el = card.select_one("h4.base-search-card__subtitle, a[class*='company']")
                    company = company_el.get_text(strip=True) if company_el else ""
                    link_el = card.select_one("a.base-card__full-link, a[class*='job-card']")
                    href = link_el.get("href", "") if link_el else ""
                    loc_el = card.select_one("span.job-search-card__location")
                    job_loc = loc_el.get_text(strip=True) if loc_el else location

                    if title:
                        jobs.append({
                            "role": title,
                            "company": company,
                            "url": href.split("?")[0] if href else "",
                            "location": job_loc,
                            "jd": "",
                            "platform": "LinkedIn",
                            "postedDaysAgo": None,
                        })
                except Exception:
                    continue
    except Exception as e:
        return jobs, str(e)

    return jobs, None


def _scrape_mcf_extended(keywords, location, max_days):
    """Second MCF search with broader keywords (e.g. 'Business Analyst') for more coverage."""
    jobs = []
    try:
        # Use related keywords to broaden results
        related_searches = [
            keywords.replace("Product Owner", "Business Analyst"),
            keywords.replace("Product Owner", "Scrum Master"),
        ]
        # Only use alternates if they differ from original
        related_searches = [s for s in related_searches if s.lower() != keywords.lower()]
        if not related_searches:
            related_searches = [f"{keywords} digital"]

        for search_term in related_searches[:2]:
            query = urllib.parse.quote_plus(search_term)
            api_url = f"https://api.mycareersfuture.gov.sg/v2/jobs?search={query}&limit=25&page=0&sortBy=new_posting_date"

            try:
                resp = http_requests.get(api_url, headers={
                    "User-Agent": "Mozilla/5.0",
                    "Accept": "application/json",
                }, timeout=15)

                if resp.status_code == 200:
                    data = resp.json()
                    for item in data.get("results", []):
                        title = item.get("title", "")
                        company = (item.get("postedCompany") or {}).get("name", "")
                        url = item.get("metadata", {}).get("jobDetailsUrl", "")
                        if not url and item.get("uuid"):
                            url = f"https://www.mycareersfuture.gov.sg/job/{item['uuid']}"

                        sal = item.get("salary", {})
                        salary_str = ""
                        if isinstance(sal, dict):
                            sal_min = sal.get("minimum", "")
                            sal_max = sal.get("maximum", "")
                            if sal_min:
                                salary_str = f"${sal_min}-${sal_max}"

                        desc_html = item.get("description", "") or ""
                        from bs4 import BeautifulSoup
                        desc_text = BeautifulSoup(desc_html, "html.parser").get_text(separator=" ", strip=True)[:4000]

                        posted = item.get("metadata", {}).get("newPostingDate", "")
                        days_ago = None
                        if posted:
                            try:
                                pd = _dt.datetime.strptime(posted[:10], "%Y-%m-%d")
                                days_ago = (_dt.datetime.now() - pd).days
                            except Exception:
                                pass

                        if title and (days_ago is None or days_ago <= max_days):
                            jobs.append({
                                "role": title,
                                "company": company,
                                "url": url,
                                "location": "Singapore",
                                "jd": desc_text,
                                "salary": salary_str,
                                "platform": "MCF+",
                                "postedDaysAgo": days_ago,
                            })
            except Exception:
                continue

    except Exception as e:
        return jobs, str(e)

    return jobs, None


def _ai_score_discovered_jobs(jobs_list):
    """Score discovered jobs with AI using active profile context. Processes in batches of 15."""
    if not jobs_list:
        return jobs_list
    
    # Only score jobs that have some useful text (title/company at minimum)
    to_score = [j for j in jobs_list if j.get("role")]
    if not to_score:
        return jobs_list

    # Get active profile for personalized scoring
    profile = get_active_profile()
    profile_summary = f"{profile.get('name', 'Candidate')} — {profile.get('headline', 'Professional')}. Skills: {', '.join(profile.get('skills', [])[:10])}. {profile.get('summary', '')[:200]}"

    BATCH_SIZE = 15
    for batch_start in range(0, len(to_score), BATCH_SIZE):
        batch = to_score[batch_start:batch_start + BATCH_SIZE]

        # Build compact job batch for AI
        batch_text = ""
        for i, j in enumerate(batch):
            jd_snippet = (j.get("jd", "") or "")[:200]
            batch_text += f"\nJOB {i+1}: {j.get('role','')} at {j.get('company','')} ({j.get('platform','')}) — {jd_snippet}\n---"

        prompt = f"""You are a career coach for the Singapore job market.

CANDIDATE PROFILE:
{profile_summary}

SCORING RULES:
- Score 1-10 based on fit with the candidate's profile
- +2 for in-house product/tech companies, -2 for consulting firms (max 4/10)
- Score 0 if "no visa sponsorship" detected
- Score 1-2 MAX for pure software engineering/developer roles (e.g. Software Engineer, Frontend Developer, Backend Engineer, DevOps, SRE) — these are NOT a match for BA/PM/PO profiles
- Score 1-2 MAX for roles requiring 8+ years when candidate has ~5 years
- Heavily penalise roles that don't match the candidate's core domain (business analysis, product management, digital transformation)
- Labels: 🔥 Strong Match (9-10), ✅ Good Fit (7-8), 🟡 Possible (5-6), ❌ Weak Fit (1-4)
- Priority: Apply Today, Apply This Week, Lower Priority, Skip

JOBS:
{batch_text}

Return ONLY a JSON array (no markdown), one object per job:
[{{"idx": 0, "score": 8, "label": "✅ Good Fit", "reason": "Two sentences.", "priority": "Apply This Week"}}]
Score every single job listed ({len(batch)} jobs). Use idx starting from 0."""

        try:
            result = call_claude(prompt)
            clean = _re.sub(r'```json|```', '', result).strip()
            rankings = json.loads(clean)
            for r in rankings:
                idx = r.get("idx", -1)
                if 0 <= idx < len(batch):
                    batch[idx]["aiScore"] = r.get("score")
                    batch[idx]["aiLabel"] = r.get("label", "")
                    batch[idx]["aiReason"] = r.get("reason", "")
                    batch[idx]["aiPriority"] = r.get("priority", "")
            print(f"[Discovery] Scored batch {batch_start//BATCH_SIZE + 1}: {len(rankings)}/{len(batch)} jobs")
        except Exception as e:
            print(f"[Discovery] AI scoring failed for batch {batch_start//BATCH_SIZE + 1}: {e}")

    # Second pass: retry any unscored jobs individually
    unscored = [j for j in to_score if j.get("aiScore") is None and j.get("role")]
    if unscored:
        print(f"[Discovery] Retrying {len(unscored)} unscored jobs individually...")
        for j in unscored[:10]:  # Limit retries
            try:
                jd_snippet = (j.get("jd", "") or "")[:300]
                mini_prompt = f"""Score this job for: {profile_summary}
Job: {j.get('role','')} at {j.get('company','')} — {jd_snippet}
Return ONLY JSON: {{"score": 7, "label": "✅ Good Fit", "reason": "Two sentences.", "priority": "Apply This Week"}}"""
                result = call_claude(mini_prompt)
                clean = _re.sub(r'```json|```', '', result).strip()
                m = _re.search(r'\{.*\}', clean, _re.DOTALL)
                if m:
                    scored = json.loads(m.group())
                    j["aiScore"] = scored.get("score")
                    j["aiLabel"] = scored.get("label", "")
                    j["aiReason"] = scored.get("reason", "")
                    j["aiPriority"] = scored.get("priority", "")
            except Exception:
                pass

    return jobs_list


@app.route("/api/discover-jobs", methods=["POST"])
def discover_jobs():
    """Search multiple job platforms and return AI-scored results."""
    data = request.json or {}
    keywords = data.get("keywords", "Product Owner")
    location = data.get("location", "Singapore")
    max_days = data.get("maxDays", 30)
    platforms = data.get("platforms", ["mycareersfuture", "linkedin_guest", "workable", "linkedin", "mcf_extended"])

    scrapers = {
        "mycareersfuture": _scrape_mycareersfuture,
        "linkedin_guest": _scrape_linkedin_guest,
        "workable": _scrape_workable,
        "linkedin": _scrape_linkedin_public,
        "mcf_extended": _scrape_mcf_extended,
    }

    all_jobs = []
    details = {}

    # Scrape platforms in parallel
    with ThreadPoolExecutor(max_workers=5) as executor:
        futures = {}
        for platform in platforms:
            if platform in scrapers:
                futures[executor.submit(scrapers[platform], keywords, location, max_days)] = platform

        for future in as_completed(futures):
            platform = futures[future]
            try:
                jobs, err = future.result(timeout=30)
                details[platform] = {"count": len(jobs), "error": err}
                all_jobs.extend(jobs)
            except Exception as e:
                details[platform] = {"count": 0, "error": str(e)}

    # Deduplicate by URL
    seen_urls = set()
    unique_jobs = []
    for j in all_jobs:
        clean_url = (j.get("url") or "").split("?")[0]
        tc = (j.get("role", "").lower() + "|" + j.get("company", "").lower())
        if clean_url and clean_url in seen_urls:
            continue
        if tc in seen_urls:
            continue
        if clean_url:
            seen_urls.add(clean_url)
        seen_urls.add(tc)
        unique_jobs.append(j)

    # Filter old jobs if we have date info
    if max_days:
        filtered = []
        for j in unique_jobs:
            if j.get("postedDaysAgo") is not None:
                if j["postedDaysAgo"] <= max_days:
                    filtered.append(j)
            else:
                # No date info — keep it
                filtered.append(j)
        unique_jobs = filtered

    # Pre-filter obviously irrelevant results (pure engineering/dev roles when searching for BA/PM)
    if unique_jobs and keywords:
        kw_lower = keywords.lower()
        is_ba_pm_search = any(t in kw_lower for t in [
            "business analyst", "product manager", "product owner", "project manager",
            "digital product", "scrum master", "ba ", "pm ", "po ",
        ])
        if is_ba_pm_search:
            irrelevant_titles = [
                "software engineer", "frontend engineer", "backend engineer",
                "full stack developer", "devops engineer", "sre ", "site reliability",
                "data engineer", "ml engineer", "machine learning engineer",
                "ios developer", "android developer", "mobile developer",
                "qa engineer", "test engineer", "security engineer",
                "infrastructure engineer", "platform engineer", "cloud engineer",
                "embedded engineer", "firmware engineer", "hardware engineer",
            ]
            before = len(unique_jobs)
            unique_jobs = [j for j in unique_jobs
                           if not any(ir in (j.get("role", "") or "").lower() for ir in irrelevant_titles)]
            removed = before - len(unique_jobs)
            if removed:
                details["relevance_filter"] = {"removed": removed, "reason": "Filtered irrelevant engineering roles"}

    # AI score the results
    if unique_jobs:
        unique_jobs = _ai_score_discovered_jobs(unique_jobs)

    # Sort by score descending
    unique_jobs.sort(key=lambda j: j.get("aiScore") or 0, reverse=True)

    return jsonify({
        "jobs": unique_jobs,
        "total": len(unique_jobs),
        "details": details
    })


# ═══════════════════════════════════════════════════════════════
# BULK AUTO-APPLY — Generate docs for multiple jobs at once
# ═══════════════════════════════════════════════════════════════

@app.route("/api/bulk-apply", methods=["POST"])
def bulk_apply():
    """
    Bulk auto-apply: for each selected job that has a JD and score >= threshold,
    generate tailored resume + cover letter .docx using AI + python-docx.
    Returns progress updates and final summary.
    """
    import base64 as b64mod

    data = request.json or {}
    job_ids = data.get("jobIds", [])
    if not job_ids:
        return jsonify({"error": "No jobs selected"}), 400

    # Load jobs from localStorage (sent from frontend)
    incoming_jobs = data.get("jobs", [])
    if not incoming_jobs:
        return jsonify({"error": "No job data provided"}), 400

    import time as _time2

    # Cap at 20 jobs per bulk run to avoid Groq rate limits (2 API calls per job)
    MAX_BULK = 20
    results = []
    generated = 0
    skipped = 0
    errors = 0
    api_call_count = 0

    for job in incoming_jobs:
        job_id = str(job.get("id", ""))
        if job_id not in [str(jid) for jid in job_ids]:
            continue

        role = job.get("role", "Unknown").strip()
        company = job.get("company", "Unknown").strip()
        jd = (job.get("jd") or "").strip()
        role_type = job.get("roleType", "Business Analyst")

        # Skip if docs already generated
        if job.get("resume_docx_b64"):
            results.append({"id": job_id, "status": "skipped", "reason": "Docs already exist"})
            skipped += 1
            continue

        # Skip if no JD
        if not jd or len(jd) < 50:
            results.append({"id": job_id, "status": "skipped", "reason": "No JD — add JD first"})
            skipped += 1
            continue

        # Cap bulk generation to avoid rate limits
        if generated >= MAX_BULK:
            results.append({"id": job_id, "status": "skipped", "reason": f"Batch limit reached ({MAX_BULK}/run). Run again to continue."})
            skipped += 1
            continue

        try:
            ai_role = is_ai_role(jd, role_type)
            P = get_active_profile()
            framing = build_product_framing(P)

            # Rate limit: 2s between each job's API calls (= ~20 req/min, under Groq limit)
            if api_call_count > 0:
                _time2.sleep(2)

            # Generate resume via AI
            resume_prompt = f"""Write a complete 2-page ATS resume for {P['name']} targeting: {role} at {company}.
JOB DESCRIPTION: {jd[:2000]}
"(AI & Personal Projects section will be added automatically)"

Use this master data — select bullets most relevant to the JD:
Contact: {P.get('mobile','')} | {P.get('email','')} | {P.get('linkedin','')}
Address: {P.get('address','')}

KPMG Singapore Feb 2021–Present bullets pool (pick 6-8):
- Defined and drove product vision, roadmap and delivery strategy for large-scale digital transformation for financial institutions
- Owned and prioritised product backlog ensuring alignment with business objectives, regulatory requirements and customer experience goals
- Drove product scope decisions through impact analysis generating ~5% additional business value
- Led sprint ceremonies (Planning, Reviews, Retros, PI Planning) across multi-squad programme, 30% velocity increase
- Managed 3rd party vendors, go-live planning, data migrations from legacy systems, 99.9% uptime
- Designed end-to-end test scenarios on Loan IQ applications, 25% defect reduction
- Supported automation and AI-enabled solution design including API integrations, saving 30 man-days
- Built business cases and executive presentations for new features securing stakeholder approval
- Collaborated with UX, architecture and engineering teams to deliver scalable digital platforms
- Drove data-informed product decisions by defining KPIs and analysing performance metrics
Key Achievements: ~5% project profit from impact analysis | 30 man-days saved via automation | Maintained delivery timeline through critical sprint-to-SIT phase

J.P. Morgan Oct 2023–Jan 2024 (pick 3): Requirements gathering from trading teams | Quantitative analysis of 5 stocks, 10% portfolio value increase | KPI measurement: Annual Return, Variance, Standard Deviation

Amazon India Mar 2018–Mar 2019 (pick 2-3): Power BI dashboards from SQL Server, 20% quality improvement | Functional/non-functional specs with 95% stakeholder satisfaction | Tableau/Power BI data visualisation, 15% efficiency increase

Skills: Agile, SAFe 6.0, JIRA, Tableau, Power BI, PSQL, Python, API integrations, Stakeholder Management, Product Vision, Roadmapping, Business Analysis, Risk Mitigation, Change Management, Budget Forecasting, KPI Tracking, AI-enabled Product Integration, SEO, GEO, CRO, Generative AI, LLM
Education: MSc Engineering Business Management, Coventry University UK (2019-2020) | BEng Electronics, Anna University India (2012-2016) | SAFe 6.0 Certified

FORMAT (plain text, no markdown):
Line 1: Amretha Karthikeyan
Line 2: #02-321 153 Gangsa Road, Singapore-670153
Line 3: Mobile: +65-90256503, email: amretha.ammu@gmail.com
Line 4: https://www.linkedin.com/in/amretha-nishanth-534b39101/
Line 5: [most relevant headline]
Then: PROFESSIONAL SUMMARY (4 sentences, JD keywords) | SKILL SET (Data visualization tools: ... / Programming: ... / Others: ... / Certification: ...) | PROFESSIONAL EXPERIENCE (company+date bold line, job title on next line, bullets) | ACADEMIC QUALIFICATION
ALL CAPS section headers. "- " bullets. No slashes on job titles. Target 750-850 words. Do NOT write HEADER. Do NOT mention the target company name anywhere in the resume."""
            resume_text = call_claude(resume_prompt)
            resume_text = _inject_ai_projects(resume_text)
            api_call_count += 1

            if resume_text.startswith("Error:") or resume_text.startswith("API error:"):
                raise Exception(f"AI error: {resume_text}")

            _time2.sleep(2)

            # Generate cover letter via AI
            cover_prompt = f"""Write a 300-word cover letter for {P['name']} applying to {role} at {company}.
{framing}
JOB DESCRIPTION: {jd[:1500]}
Plain text, specific to company, include metrics."""
            cover_text = call_claude(cover_prompt)
            api_call_count += 1

            if cover_text.startswith("Error:") or cover_text.startswith("API error:"):
                raise Exception(f"AI error: {cover_text}")

            # Create .docx
            resume_bytes = _create_docx_from_text(resume_text)
            cover_bytes = _create_docx_from_text(cover_text)

            results.append({
                "id": job_id,
                "status": "generated",
                "resume_docx_b64": b64mod.b64encode(resume_bytes).decode(),
                "cover_docx_b64": b64mod.b64encode(cover_bytes).decode(),
                "resume_variant": "AI",
                "resume_filename": f"Resume_{company.replace(' ','_')}.docx",
                "cover_filename": f"CoverLetter_{company.replace(' ','_')}.docx",
            })
            generated += 1
        except Exception as e:
            results.append({"id": job_id, "status": "error", "reason": str(e)[:100]})
            errors += 1

    return jsonify({
        "results": results,
        "generated": generated,
        "skipped": skipped,
        "errors": errors,
        "total": len(job_ids)
    })


# ═══════════════════════════════════════════════════════════════
# AGENT SYSTEM
# ═══════════════════════════════════════════════════════════════

import threading
import datetime

AGENT_CRON_SECRET    = os.environ.get("AGENT_CRON_SECRET", "jobhunt2025")
NOTIFICATION_PHONE   = os.environ.get("NOTIFICATION_PHONE", "")   # e.g. whatsapp:+6590256503
TWILIO_ACCOUNT_SID   = os.environ.get("TWILIO_ACCOUNT_SID", "")
TWILIO_AUTH_TOKEN    = os.environ.get("TWILIO_AUTH_TOKEN", "")
TWILIO_WHATSAPP_FROM = os.environ.get("TWILIO_WHATSAPP_FROM", "whatsapp:+14155238886")  # Twilio sandbox

def _get_twilio_sid():   return os.environ.get("TWILIO_ACCOUNT_SID", "") or get_setting("TWILIO_ACCOUNT_SID") or ""
def _get_twilio_token(): return os.environ.get("TWILIO_AUTH_TOKEN", "") or get_setting("TWILIO_AUTH_TOKEN") or ""
def _get_whatsapp_to():  return os.environ.get("NOTIFICATION_PHONE", "") or get_setting("NOTIFICATION_PHONE")


def send_whatsapp(message):
    """Send a WhatsApp message via Twilio API."""
    sid   = _get_twilio_sid()
    token = _get_twilio_token()
    to    = _get_whatsapp_to()
    frm   = TWILIO_WHATSAPP_FROM
    if not sid or not token or not to:
        print("WhatsApp not configured — set TWILIO_ACCOUNT_SID, TWILIO_AUTH_TOKEN, NOTIFICATION_PHONE")
        return False
    try:
        # Ensure to/from are prefixed with whatsapp:
        if not to.startswith("whatsapp:"): to = f"whatsapp:{to}"
        if not frm.startswith("whatsapp:"): frm = f"whatsapp:{frm}"
        url  = f"https://api.twilio.com/2010-04-01/Accounts/{sid}/Messages.json"
        resp = http_requests.post(url, auth=(sid, token), data={"From": frm, "To": to, "Body": message})
        if resp.status_code in (200, 201):
            print(f"[WhatsApp] Sent: {message[:60]}...")
            return True
        else:
            print(f"[WhatsApp] Failed {resp.status_code}: {resp.text[:200]}")
            return False
    except Exception as e:
        print(f"[WhatsApp] Error: {e}")
        return False

# Keep send_email as alias so existing calls still work
def send_email(subject, html_body):
    # Strip HTML tags for WhatsApp plain text
    import re as _re
    text = _re.sub(r'<[^>]+>', '', html_body)
    text = _re.sub(r'\s+', ' ', text).strip()
    msg  = f"*{subject}*\n\n{text[:1000]}"
    return send_whatsapp(msg)


def _send_whatsapp_summary(summary):
    """Send a concise, well-formatted WhatsApp summary after agent run."""
    top = summary.get("top_jobs", [])
    labels = {"import": "Import", "cron": "Daily 9AM", "manual": "Manual"}
    trigger = labels.get(summary.get("trigger", ""), "Run")
    now_str = datetime.datetime.now().strftime("%d %b %H:%M")

    msg = f"🤖 *Job Agent Complete* — {trigger} · {now_str}\n\n"
    msg += f"📊 *{summary['total']}* processed · *{summary['scored']}* scored · *{summary['docs']}* docs ready\n\n"

    if top:
        msg += "🏆 *Top Matches:*\n"
        for j in top[:5]:
            score = j.get("aiScore", "?")
            company = j.get("company", "")
            role = j.get("role", "")
            priority = j.get("aiPriority", "")
            has_docs = "📄" if j.get("resume_docx_b64") else ""
            msg += f"• {score}/10 — {company} · {role} {has_docs}\n"
            if priority:
                msg += f"  → _{priority}_\n"
    else:
        msg += "No scored jobs yet.\n"

    msg += "\n🔗 Open tracker: https://job-hunt-app-r7my.onrender.com"
    return send_whatsapp(msg)


def agent_process_job(job):
    """Full agent pipeline for one job: score → generate docs → save."""
    log     = []
    job_id  = job.get("id")
    role    = job.get("role", "Unknown")
    company = job.get("company", "Unknown")
    jd      = job.get("jd", "") or ""

    log.append(f"Processing: {role} @ {company}")

    # STEP 1: AI Score
    if jd and len(jd) > 50 and job.get("aiScore") is None:
        try:
            P = get_active_profile()
            skills_short = ', '.join(P.get('skills', [])[:8])
            prompt = f"""You are a career coach for the tech job market.

CANDIDATE: {P.get('name','Unknown')} — {P.get('headline','')}
Summary: {P.get('summary','')[:200]}
Skills: {skills_short}
Certification: {P.get('certification','')}
Target: In-house product roles (NOT consulting).

SCORING RULES:
- Score 1-10 based on fit
- +2 for in-house product companies (Grab, Sea, Airwallex, Stripe, GovTech, startups)
- -2 for consulting (KPMG, Deloitte, Accenture, Big4) — max 4/10 for consulting roles
- Score 0 if JD says "no visa sponsorship"
- Labels: 🔥 Strong Match (9-10), ✅ Good Fit (7-8), 🟡 Possible (5-6), ❌ Weak Fit (1-4)
- Priority: Apply Today, Apply This Week, Lower Priority, Skip

JOB:
Title: {role}
Company: {company}
JD: {jd[:500]}

Return ONLY valid JSON, no markdown:
{{"score": 8, "label": "✅ Good Fit", "reason": "Two sentence reason.", "priority": "Apply This Week"}}"""

            result = call_claude(prompt)
            clean  = result.strip().strip("```json").strip("```").strip()
            # Find JSON object in response
            import re
            m = re.search(r'\{.*\}', clean, re.DOTALL)
            if m:
                scored = json.loads(m.group())
                job["aiScore"]    = scored.get("score")
                job["aiLabel"]    = scored.get("label", "")
                job["aiReason"]   = scored.get("reason", "")
                job["aiPriority"] = scored.get("priority", "")
                log.append(f"  Scored: {job['aiLabel']} ({job['aiScore']}/10)")
        except Exception as e:
            log.append(f"  Scoring failed: {e}")
    elif job.get("aiScore") is not None:
        log.append(f"  Already scored: {job.get('aiLabel')} ({job.get('aiScore')}/10)")
    else:
        log.append(f"  No JD — skipping score")

    # STEP 2: Generate docs (score >= 5, no docs yet)
    score = job.get("aiScore") or 0
    if score >= 5 and not job.get("resume_docx_b64"):
        try:
            import base64 as b64mod
            ai_role = is_ai_role(jd, job.get("roleType", ""))
            P_agent = get_active_profile()
            framing_agent = build_product_framing(P_agent)

            # Build structured experience text for resume
            exp_lines_r = ""
            for exp in P_agent.get('experience', []):
                exp_lines_r += f"\n--- {exp.get('role','')} at {exp.get('company','')} ({exp.get('period','')}) ---\n"
                for b in exp.get('bullets', []):
                    exp_lines_r += f"  - {b}\n"
            skills_r = ', '.join(P_agent.get('skills', []))

            resume_prompt = f"""Write a COMPLETE ATS-optimised resume (750+ words, plain text, no markdown) for {P_agent['name']} targeting: {role} at {company}.
{framing_agent}
JOB DESCRIPTION: {jd[:2000]}
{"AI ROLE: Feature AI projects with URL: " + P_agent.get('aiProjectUrl','') if ai_role else ""}
CANDIDATE: {P_agent.get('mobile','')}, {P_agent.get('email','')}, {P_agent.get('linkedin','')}
Skills: {skills_r}
Experience:
{exp_lines_r}
Education: {'; '.join([e.get('degree','') + ' — ' + e.get('school','') for e in P_agent.get('education',[])])}
Certification: {P_agent.get('certification','')}

Include ALL {len(P_agent.get('experience',[]))} roles with 5-8 bullets each. Include 25+ skills. Include education + SAFe cert.
ALL CAPS headers. "- " for bullets. Every bullet = verb + metric. At least 750 words."""
            resume_text = call_claude(resume_prompt, max_tokens=8192)

            cover_prompt = f"""Write a 300-350 word cover letter for {P_agent['name']} applying to {role} at {company}.
{framing_agent}
JOB DESCRIPTION: {jd[:2000]}
{"AI ROLE: Mention AI project with URL." if ai_role else ""}
Plain text, specific to company, include metrics (~5% business value, 30 man-days). Reference SAFe certification."""
            cover_text = call_claude(cover_prompt)

            resume_bytes = _create_docx_from_text(resume_text)
            cover_bytes = _create_docx_from_text(cover_text)

            job["resume_docx_b64"]     = b64mod.b64encode(resume_bytes).decode()
            job["cover_docx_b64"]      = b64mod.b64encode(cover_bytes).decode()
            job["resume_variant"]      = "AI"
            job["resume_filename"]     = f"Resume_{company.replace(' ','_')}.docx"
            job["cover_filename"]      = f"CoverLetter_{company.replace(' ','_')}.docx"
            job["resume_generated_at"] = datetime.datetime.utcnow().isoformat()
            log.append(f"  Docs generated (AI + python-docx)")
        except Exception as e:
            log.append(f"  Doc gen error: {e}")
    elif score < 5:
        log.append(f"  Score {score}/10 — skipping docs")
    else:
        log.append(f"  Docs already exist")

    # STEP 3: Save to Supabase
    try:
        sb = get_supabase()
        if sb:
            row = {
                "id":                  str(job.get("id", "")),
                "linkedInId":          job.get("linkedInId", ""),
                "role":                role,
                "company":             company,
                "status":              job.get("status", "saved"),
                "url":                 job.get("url", ""),
                "jd":                  jd[:8000],
                "roleType":            job.get("roleType", ""),
                "source":              job.get("source", ""),
                "salary":              job.get("salary", ""),
                "dateApplied":         job.get("dateApplied", ""),
                "aiScore":             job.get("aiScore"),
                "aiLabel":             job.get("aiLabel", ""),
                "aiReason":            job.get("aiReason", ""),
                "aiPriority":          job.get("aiPriority", ""),
                "notes":               job.get("notes", ""),
                "resume_docx_b64":     (job.get("resume_docx_b64") or "")[:500000],
                "cover_docx_b64":      (job.get("cover_docx_b64") or "")[:500000],
                "resume_variant":      job.get("resume_variant", ""),
                "resume_filename":     job.get("resume_filename", ""),
                "cover_filename":      job.get("cover_filename", ""),
                "resume_generated_at": job.get("resume_generated_at", ""),
            }
            sb.table("jobs").upsert(row, on_conflict="id").execute()
            log.append(f"  Saved to Supabase")
    except Exception as e:
        log.append(f"  Supabase save failed: {e}")

    return job, log


def agent_run(jobs_to_process, trigger="manual"):
    """Run agent pipeline over list of jobs, then notify."""
    results  = []
    all_logs = []
    scored   = []
    docs_gen = []

    for job in jobs_to_process:
        enriched, log = agent_process_job(job)
        all_logs.extend(log)
        results.append(enriched)
        if enriched.get("aiScore") is not None:
            scored.append(enriched)
        if enriched.get("resume_docx_b64"):
            docs_gen.append(enriched)

    top_jobs = sorted(scored, key=lambda j: j.get("aiScore", 0), reverse=True)[:5]
    summary  = {
        "trigger":  trigger,
        "total":    len(jobs_to_process),
        "scored":   len(scored),
        "docs":     len(docs_gen),
        "top_jobs": top_jobs,
        "logs":     all_logs,
        "results":  results,
    }
    _send_agent_notifications(summary)
    return summary


def _send_agent_notifications(summary):
    """Send email notification after agent run."""
    top     = summary["top_jobs"]
    labels  = {"import": "📥 Auto (import)", "cron": "⏰ Daily", "manual": "▶ Manual"}
    trigger_label = labels.get(summary["trigger"], "▶ Run")
    now_str = datetime.datetime.now().strftime("%d %b %Y %H:%M")

    # Email only
    rows_html = ""
    for j in top:
        has_docs = "✅ Ready" if j.get("resume_docx_b64") else "—"
        rows_html += (
            f"<tr>"
            f"<td style='padding:10px;border-bottom:1px solid #e5e7eb;'>"
            f"<a href='{j.get('url','#')}' style='font-weight:700;color:#1d4ed8;'>{j.get('company','')}</a><br>"
            f"<span style='font-size:13px;color:#374151;'>{j.get('role','')}</span></td>"
            f"<td style='padding:10px;border-bottom:1px solid #e5e7eb;text-align:center;font-weight:700;'>{j.get('aiScore','?')}/10</td>"
            f"<td style='padding:10px;border-bottom:1px solid #e5e7eb;'>{j.get('aiLabel','')}</td>"
            f"<td style='padding:10px;border-bottom:1px solid #e5e7eb;font-size:12px;color:#6b7280;'>{j.get('aiReason','')}</td>"
            f"<td style='padding:10px;border-bottom:1px solid #e5e7eb;text-align:center;'>{has_docs}</td>"
            f"</tr>"
        )

    html = f"""<html><body style="font-family:-apple-system,sans-serif;background:#f9fafb;padding:20px;">
<div style="max-width:720px;margin:0 auto;background:white;border-radius:12px;overflow:hidden;box-shadow:0 1px 3px rgba(0,0,0,.1);">
  <div style="background:linear-gradient(135deg,#1d4ed8,#7c3aed);padding:28px 32px;color:white;">
    <h1 style="margin:0;font-size:22px;">🤖 Job Agent Complete</h1>
    <p style="margin:6px 0 0;opacity:.85;">{trigger_label} · {now_str}</p>
  </div>
  <div style="padding:24px 32px;">
    <div style="display:flex;gap:16px;margin-bottom:24px;">
      <div style="flex:1;background:#eff6ff;border-radius:8px;padding:14px;text-align:center;">
        <div style="font-size:28px;font-weight:700;color:#1d4ed8;">{summary['total']}</div>
        <div style="font-size:12px;color:#6b7280;">Processed</div>
      </div>
      <div style="flex:1;background:#f0fdf4;border-radius:8px;padding:14px;text-align:center;">
        <div style="font-size:28px;font-weight:700;color:#16a34a;">{summary['scored']}</div>
        <div style="font-size:12px;color:#6b7280;">Scored</div>
      </div>
      <div style="flex:1;background:#faf5ff;border-radius:8px;padding:14px;text-align:center;">
        <div style="font-size:28px;font-weight:700;color:#7c3aed;">{summary['docs']}</div>
        <div style="font-size:12px;color:#6b7280;">Docs Ready</div>
      </div>
    </div>
    {'<h2 style="font-size:16px;margin-bottom:12px;">🏆 Top Matches</h2><table style="width:100%;border-collapse:collapse;"><thead><tr style="background:#f3f4f6;"><th style="padding:10px;text-align:left;font-size:12px;color:#6b7280;">JOB</th><th style="padding:10px;font-size:12px;color:#6b7280;">SCORE</th><th style="padding:10px;font-size:12px;color:#6b7280;">FIT</th><th style="padding:10px;font-size:12px;color:#6b7280;">REASON</th><th style="padding:10px;font-size:12px;color:#6b7280;">DOCS</th></tr></thead><tbody>' + rows_html + '</tbody></table>' if top else '<p style="color:#6b7280;">No scored jobs yet.</p>'}
  </div>
  <div style="background:#f3f4f6;padding:16px 32px;text-align:center;font-size:13px;color:#6b7280;">
    Open your <a href="https://job-hunt-app.onrender.com" style="color:#1d4ed8;">Job Tracker</a> to download documents
  </div>
</div></body></html>"""

    send_email(
        subject=f"🤖 Agent: {summary['scored']} scored, {summary['docs']} docs ready — {datetime.datetime.now().strftime('%d %b')}",
        html_body=html
    )

    # Also send concise WhatsApp summary
    _send_whatsapp_summary(summary)


# ─── AUTONOMOUS AGENT PIPELINE ───────────────────────────────
def agent_autonomous_pipeline(config=None):
    """
    Fully autonomous agentic workflow:
      Step 1: Discover new jobs from MCF/LinkedIn/Workable
      Step 2: Scrape LinkedIn saved jobs (if credentials available)
      Step 3: Merge & deduplicate all discovered jobs into Supabase
      Step 4: AI-score all unscored jobs
      Step 5: Generate resume + cover letter for top-scoring jobs
      Step 6: Save everything to Supabase
      Step 7: Send WhatsApp + email notification summary
    """
    config = config or {}
    pipeline_log = []
    P = get_active_profile()

    def log(msg):
        pipeline_log.append(msg)
        print(f"[Agent] {msg}")

    log(f"🤖 Starting autonomous pipeline for {P.get('name', 'user')}")

    # ── Step 1: Job Discovery from web scrapers ──
    discovered_jobs = []
    keywords = config.get("keywords", P.get("headline", "Product Manager"))
    location = config.get("location", "Singapore")
    max_days = config.get("max_days", 30)
    platforms = config.get("platforms", ["mycareersfuture", "linkedin_guest", "workable"])

    if platforms:
        from concurrent.futures import ThreadPoolExecutor
        scraper_map = {
            "mycareersfuture": _scrape_mycareersfuture,
            "linkedin_guest": _scrape_linkedin_guest,
            "workable": _scrape_workable,
            "mcf_extended": _scrape_mcf_extended,
            "linkedin": _scrape_linkedin_public,
        }
        log(f"Step 1: Discovering jobs — keywords='{keywords}', location='{location}', platforms={platforms}")
        with ThreadPoolExecutor(max_workers=3) as pool:
            futures = {}
            for p_name in platforms:
                fn = scraper_map.get(p_name)
                if fn:
                    futures[p_name] = pool.submit(fn, keywords, location, max_days)
            for p_name, fut in futures.items():
                try:
                    results = fut.result(timeout=60)
                    discovered_jobs.extend(results)
                    log(f"  {p_name}: {len(results)} jobs found")
                except Exception as e:
                    log(f"  {p_name}: error — {str(e)[:60]}")

        # Deduplicate by title+company
        seen = set()
        unique = []
        for j in discovered_jobs:
            key = f"{j.get('role','').lower().strip()}|{j.get('company','').lower().strip()}"
            if key not in seen:
                seen.add(key)
                unique.append(j)
        discovered_jobs = unique
        log(f"  Total unique discovered: {len(discovered_jobs)}")
    else:
        log("Step 1: Skipping discovery (no platforms configured)")

    # ── Step 2: (LinkedIn Selenium scrape removed — using discovery scrapers only) ──
    linkedin_jobs = []
    log("Step 2: Skipping LinkedIn login scrape (removed — use bookmarklet instead)")

    # ── Step 3: Merge & sync to Supabase ──
    all_new_jobs = discovered_jobs + linkedin_jobs
    total_added = 0
    total_skipped = 0

    if all_new_jobs:
        log(f"Step 3: Syncing {len(all_new_jobs)} jobs to Supabase...")

        # Sync discovered jobs
        if discovered_jobs:
            sb = get_supabase()
            if sb:
                try:
                    existing = sb.table("jobs").select("id,url").execute().data or []
                    existing_urls = {(j.get("url") or "").split("?")[0] for j in existing if j.get("url")}
                    import time as _time
                    to_insert = []
                    for dj in discovered_jobs:
                        clean_url = (dj.get("url") or "").split("?")[0]
                        if clean_url and clean_url in existing_urls:
                            total_skipped += 1
                            continue
                        job_id = str(int(_time.time() * 1000)) + str(len(to_insert))
                        to_insert.append({
                            "id": job_id,
                            "role": dj.get("role", ""),
                            "company": dj.get("company", ""),
                            "url": clean_url,
                            "jd": (dj.get("jd") or "")[:8000],
                            "status": "saved",
                            "source": dj.get("source", "Discovery"),
                            "roleType": "Business Analyst",
                            "dateApplied": datetime.datetime.now().isoformat(),
                        })
                        if clean_url:
                            existing_urls.add(clean_url)
                    if to_insert:
                        sb.table("jobs").upsert(to_insert, on_conflict="id").execute()
                    total_added += len(to_insert)
                    log(f"  Discovery sync: {len(to_insert)} new, {total_skipped} duplicates")
                except Exception as e:
                    log(f"  Discovery sync error: {str(e)[:60]}")
    else:
        log("Step 3: No new jobs to sync")

    log(f"  Summary: {total_added} added, {total_skipped} skipped")

    # ── Step 4 + 5 + 6: Score, generate docs, save (via existing agent_run) ──
    sb = get_supabase()
    agent_summary = None
    if sb:
        try:
            res = sb.table("jobs").select("*").execute()
            jobs = [j for j in (res.data or []) if not j.get("isDemo")]
            to_run = [
                j for j in jobs
                if (j.get("jd") and j.get("aiScore") is None) or
                   (j.get("aiScore", 0) >= 5 and not j.get("resume_docx_b64"))
            ]
            if to_run:
                log(f"Step 4-6: Processing {len(to_run)} jobs (score → docs → save)...")
                agent_summary = agent_run(to_run, trigger="auto")
                log(f"  Done: {agent_summary['scored']} scored, {agent_summary['docs']} docs generated")
            else:
                log("Step 4-6: All jobs already processed")
        except Exception as e:
            log(f"Step 4-6 error: {str(e)[:60]}")

    log("✅ Autonomous pipeline complete")

    return {
        "pipeline_log": pipeline_log,
        "discovered": len(discovered_jobs),
        "linkedin": len(linkedin_jobs),
        "added": total_added,
        "skipped": total_skipped,
        "scored": agent_summary["scored"] if agent_summary else 0,
        "docs": agent_summary["docs"] if agent_summary else 0,
    }


@app.route("/api/agent/autonomous", methods=["POST"])
def agent_autonomous_route():
    """Trigger the fully autonomous agentic pipeline."""
    data = request.json or {}
    config = {
        "keywords": data.get("keywords", ""),
        "location": data.get("location", "Singapore"),
        "max_days": data.get("max_days", 30),
        "platforms": data.get("platforms", ["mycareersfuture", "linkedin_guest", "workable"]),
    }

    def bg():
        with app.app_context():
            agent_autonomous_pipeline(config)

    threading.Thread(target=bg, daemon=True).start()
    return jsonify({"status": "started", "message": "Autonomous agent pipeline running in background"})


# ─── AGENT ROUTES ────────────────────────────────────────────

@app.route("/api/agent/run", methods=["POST"])
def agent_run_route():
    """Manual Run Agent — processes all pending jobs in background."""
    data      = request.json or {}
    force_all = data.get("force_all", False)

    sb = get_supabase()
    if not sb:
        return jsonify({"error": "Supabase not configured"}), 400
    try:
        res      = sb.table("jobs").select("*").execute()
        all_jobs = res.data or []
    except Exception as e:
        return jsonify({"error": str(e)}), 500

    if force_all:
        to_run = [j for j in all_jobs if not j.get("isDemo")]
    else:
        to_run = [
            j for j in all_jobs
            if not j.get("isDemo") and (
                (j.get("jd") and j.get("aiScore") is None) or
                (j.get("aiScore", 0) >= 5 and not j.get("resume_docx_b64"))
            )
        ]

    if not to_run:
        return jsonify({"status": "nothing_to_do", "message": "All jobs already processed"})

    def bg():
        with app.app_context():
            agent_run(to_run, trigger="manual")

    threading.Thread(target=bg, daemon=True).start()
    return jsonify({"status": "started", "count": len(to_run),
                    "message": f"Agent processing {len(to_run)} jobs in background"})


@app.route("/api/agent/run-import", methods=["POST"])
def agent_run_import():
    """Auto-triggered after bookmarklet import — processes new jobs immediately."""
    data     = request.json or {}
    new_jobs = data.get("jobs", [])
    if not new_jobs:
        return jsonify({"status": "nothing_to_do"})

    def bg():
        with app.app_context():
            agent_run(new_jobs, trigger="import")

    threading.Thread(target=bg, daemon=True).start()
    return jsonify({"status": "started", "count": len(new_jobs)})


@app.route("/api/config/save", methods=["POST"])
def config_save():
    """Save LinkedIn + email credentials to Supabase config table."""
    sb = get_supabase()
    if not sb:
        return jsonify({"error": "Supabase not configured"}), 400
    data = request.json or {}
    try:
        # Store each key as a row in a simple config table
        for key, val in data.items():
            if val:  # only save non-empty values
                sb.table("config").upsert({"key": key, "value": val}).execute()
        return jsonify({"ok": True})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/config/load", methods=["GET"])
def config_load():
    """Load config from Supabase. Returns keys without sensitive values (masked)."""
    sb = get_supabase()
    if not sb:
        return jsonify({"error": "Supabase not configured"}), 400
    try:
        rows = sb.table("config").select("key,value").execute().data or []
        result = {}
        for row in rows:
            k, v = row["key"], row.get("value", "")
            # Mask passwords/sensitive values
            if "password" in k.lower() or "secret" in k.lower():
                result[k] = "••••••••" if v else ""
            else:
                result[k] = v
        return jsonify(result)
    except Exception as e:
        return jsonify({"error": str(e)}), 500


def get_config_value(key):
    """Retrieve a single config value from Supabase config table."""
    # First check env vars (env vars take priority)
    env_map = {
        "linkedin_email":    "LINKEDIN_EMAIL",
        "linkedin_password": "LINKEDIN_PASSWORD",
        "twilio_account_sid": "TWILIO_ACCOUNT_SID",
        "twilio_auth_token":  "TWILIO_AUTH_TOKEN",
        "notification_phone": "NOTIFICATION_PHONE",
    }
    env_key = env_map.get(key)
    if env_key:
        val = os.environ.get(env_key, "")
        if val:
            return val
    # Fall back to Supabase config table
    sb = get_supabase()
    if not sb:
        return ""
    try:
        rows = sb.table("config").select("value").eq("key", key).execute().data or []
        return rows[0]["value"] if rows else ""
    except Exception:
        return ""


@app.route("/api/agent/status", methods=["GET"])
def agent_status():
    """Return counts of processed vs pending jobs."""
    empty = {"total": 0, "with_jd": 0, "scored": 0, "with_docs": 0, "pending": 0}
    sb = get_supabase()
    if not sb:
        return jsonify(empty)
    try:
        res  = sb.table("jobs").select("id,aiScore,resume_docx_b64,jd").execute()
        jobs = res.data or []
        return jsonify({
            "total":     len(jobs),
            "with_jd":   sum(1 for j in jobs if (j.get("jd") or "").strip()),
            "scored":    sum(1 for j in jobs if j.get("aiScore") is not None),
            "with_docs": sum(1 for j in jobs if j.get("resume_docx_b64")),
            "pending":   sum(1 for j in jobs if j.get("aiScore") is None),
        })
    except Exception as e:
        return jsonify(empty)



# ─── CREDENTIALS STORE ──────────────────────────────────────────────────────
# Stores LinkedIn + Gmail credentials in Supabase settings table
# so user can enter them via the UI without touching Render env vars.

def get_setting(key):
    """Read a setting from env vars, Supabase, or memory cache."""
    # First try env var (Render dashboard)
    env_val = os.environ.get(key.upper(), "")
    if env_val:
        return env_val
    # Then try in-memory cache
    if key in _settings_cache:
        return _settings_cache[key]
    # Then try Supabase settings table
    try:
        sb = get_supabase()
        if not sb:
            return ""
        if ensure_settings_table():
            res = sb.table("settings").select("value").eq("key", key).execute()
            if res.data:
                val = res.data[0]["value"]
                _settings_cache[key] = val  # Cache for future reads
                return val
    except Exception:
        pass
    return ""

_settings_table_ok = False
_settings_cache = {}  # In-memory fallback cache

def ensure_settings_table():
    """Check if settings table exists. If not, mark as unavailable and use memory/env fallback."""
    global _settings_table_ok
    if _settings_table_ok:
        return True
    sb = get_supabase()
    if not sb:
        return False
    try:
        sb.table("settings").select("key").limit(1).execute()
        _settings_table_ok = True
        return True
    except Exception:
        # Table doesn't exist — try SQL creation via Supabase REST SQL endpoint
        try:
            resp = http_requests.post(
                f"{SUPABASE_URL}/rest/v1/rpc/exec_sql",
                headers={
                    "apikey": SUPABASE_KEY,
                    "Authorization": f"Bearer {SUPABASE_KEY}",
                    "Content-Type": "application/json"
                },
                json={"query": "CREATE TABLE IF NOT EXISTS public.settings (key TEXT PRIMARY KEY, value TEXT NOT NULL, updated_at TIMESTAMPTZ DEFAULT NOW()); ALTER TABLE public.settings ENABLE ROW LEVEL SECURITY; CREATE POLICY IF NOT EXISTS settings_all ON public.settings FOR ALL USING (true);"},
                timeout=10
            )
            if resp.status_code < 300:
                _settings_table_ok = True
                print("[Settings] Created settings table via REST RPC")
                return True
        except Exception:
            pass
        # If we reach here, just use in-memory cache — don't crash
        print("[Settings] Table unavailable — using in-memory + env fallback")
        return False


def upsert_setting(key, value):
    """Save a setting to Supabase settings table, with memory fallback."""
    global _settings_cache
    _settings_cache[key] = value  # Always cache in memory as backup
    try:
        sb = get_supabase()
        if not sb:
            return True  # Saved in memory
        if ensure_settings_table():
            sb.table("settings").upsert({"key": key, "value": value}, on_conflict="key").execute()
            print(f"[Settings] Saved {key} to Supabase")
            return True
        else:
            print(f"[Settings] Saved {key} to memory (Supabase unavailable)")
            return True  # Memory fallback succeeded
    except Exception as e:
        print(f"[Settings] upsert for {key} failed in Supabase, kept in memory: {e}")
        return True  # Memory fallback succeeded

@app.route("/api/settings/save", methods=["POST"])
def save_settings():
    data   = request.json or {}
    secret = data.get("secret", "")
    if secret != AGENT_CRON_SECRET:
        return jsonify({"error": "Unauthorized"}), 401
    saved  = []
    failed = []
    for key in ["LINKEDIN_EMAIL", "LINKEDIN_PASSWORD", "LI_AT_COOKIE", "TWILIO_ACCOUNT_SID", "TWILIO_AUTH_TOKEN", "NOTIFICATION_PHONE"]:
        val = data.get(key, "").strip()
        if val:
            ok = upsert_setting(key, val)
            (saved if ok else failed).append(key)
    if failed:
        return jsonify({"status": "partial", "saved": saved, "failed": failed,
                        "error": f"Could not save {failed} — make sure the settings table exists in Supabase (run supabase_setup.sql)"})
    return jsonify({"status": "ok", "saved": saved})

@app.route("/api/settings/load", methods=["GET"])
def load_settings():
    """Return non-sensitive settings (mask passwords)."""
    return jsonify({
        "LINKEDIN_EMAIL":     _get_linkedin_email(),
        "TWILIO_ACCOUNT_SID": _get_twilio_sid(),
        "NOTIFICATION_PHONE": _get_whatsapp_to(),
        "linkedin_pw_set":    bool(_get_linkedin_password()),
        "li_at_set":          bool(_get_li_at_cookie()),
        "twilio_token_set":   bool(_get_twilio_token()),
    })


# ─── LINKEDIN SAVED JOBS (credentials + cookie, no Selenium) ─────────────────

def _get_linkedin_email():    return get_setting("LINKEDIN_EMAIL") or os.environ.get("LINKEDIN_EMAIL", "")
def _get_linkedin_password(): return get_setting("LINKEDIN_PASSWORD") or os.environ.get("LINKEDIN_PASSWORD", "")
def _get_li_at_cookie():      return get_setting("LI_AT_COOKIE") or os.environ.get("LI_AT_COOKIE", "")


def _linkedin_login_for_cookie():
    """
    Log in to LinkedIn with email + password via HTTP and return the li_at cookie.
    No Selenium / Chrome needed — pure requests.
    Returns the li_at cookie string, or None on failure.
    """
    from bs4 import BeautifulSoup
    import time as _t

    email = _get_linkedin_email()
    pw    = _get_linkedin_password()
    if not email or not pw:
        return None

    sess = http_requests.Session()
    sess.headers.update({
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
                       "AppleWebKit/537.36 (KHTML, like Gecko) Chrome/122.0.0.0 Safari/537.36",
        "Accept-Language": "en-US,en;q=0.9",
    })

    # Step 1: GET the login page to grab the CSRF token
    try:
        print("[LinkedIn Login] Fetching login page...")
        r = sess.get("https://www.linkedin.com/login", timeout=20)
        if r.status_code != 200:
            print(f"[LinkedIn Login] Login page returned {r.status_code}")
            return None
    except Exception as e:
        print(f"[LinkedIn Login] Error fetching login page: {e}")
        return None

    soup = BeautifulSoup(r.text, "html.parser")
    csrf_field = soup.find("input", {"name": "loginCsrfParam"})
    if not csrf_field:
        print("[LinkedIn Login] Could not find loginCsrfParam on login page")
        return None
    csrf_val = csrf_field.get("value", "")
    print(f"[LinkedIn Login] Got CSRF token, posting credentials...")

    _t.sleep(1)  # polite delay to look human

    # Step 2: POST credentials
    login_data = {
        "session_key": email,
        "session_password": pw,
        "loginCsrfParam": csrf_val,
        "trk": "guest_homepage-basic_sign-in-submit",
    }
    try:
        r2 = sess.post(
            "https://www.linkedin.com/checkpoint/lg/login-submit",
            data=login_data,
            timeout=20,
            allow_redirects=True,
        )
    except Exception as e:
        print(f"[LinkedIn Login] Error posting credentials: {e}")
        return None

    # Step 3: Check for li_at cookie in the session
    li_at = sess.cookies.get("li_at", domain=".linkedin.com") or sess.cookies.get("li_at")
    if not li_at:
        # Sometimes the cookie is in the response headers
        for cookie in sess.cookies:
            if cookie.name == "li_at":
                li_at = cookie.value
                break

    if li_at:
        print(f"[LinkedIn Login] Success — obtained li_at cookie ({len(li_at)} chars)")
        # Cache the cookie in settings so we don't re-login every time
        try:
            upsert_setting("LI_AT_COOKIE", li_at)
            print("[LinkedIn Login] Cached li_at cookie in settings for reuse")
        except Exception:
            pass
        return li_at

    # Check if we hit a security challenge
    if "challenge" in r2.url or "checkpoint" in r2.url:
        print(f"[LinkedIn Login] Security challenge detected: {r2.url}")
        print("[LinkedIn Login] You may need to verify from your browser first, then retry")
    else:
        print(f"[LinkedIn Login] Login did not return li_at cookie. Final URL: {r2.url}")

    return None


def _get_or_login_li_at():
    """
    Get the li_at cookie — try stored cookie first, then auto-login with credentials.
    Returns (li_at_string, source_label) or (None, error_label).
    """
    # 1. Check for a stored / manually-provided li_at cookie
    li_at = _get_li_at_cookie()
    if li_at:
        li_at = li_at.strip().strip('"').strip("'")
        return li_at, "stored_cookie"

    # 2. Try credential-based login
    print("[LinkedIn] No stored li_at cookie — attempting credential login...")
    li_at = _linkedin_login_for_cookie()
    if li_at:
        return li_at, "credential_login"

    return None, "no_credentials"


def linkedin_scrape_saved_jobs_via_cookie(max_days=30):
    """
    Scrape LinkedIn saved jobs using the li_at session cookie.
    Tries stored cookie first, falls back to email/password login.
    Uses LinkedIn's internal Voyager REST API — lightweight HTTP requests only.
    Returns (jobs_list, error_message).
    """
    import re, uuid, time as _time

    li_at, source = _get_or_login_li_at()
    if not li_at:
        email = _get_linkedin_email()
        if email:
            return [], "Login failed — LinkedIn may require a security verification. Log in from your browser, then retry."
        return [], "no_credentials"

    # Clean the cookie value (user might copy with quotes)
    li_at = li_at.strip().strip('"').strip("'")

    csrf_token = f"ajax:{uuid.uuid4()}"
    headers = {
        "User-Agent": "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) "
                      "AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36",
        "Accept": "application/vnd.linkedin.normalized+json+2.1",
        "x-restli-protocol-version": "2.0.0",
        "csrf-token": csrf_token,
    }
    cookies = {
        "li_at": li_at,
        "JSESSIONID": f'"{ csrf_token }"',
    }

    all_jobs = []
    start = 0
    page_size = 40
    max_pages = 15  # safety limit
    cutoff = datetime.datetime.now(datetime.timezone.utc) - datetime.timedelta(days=max_days)

    for page in range(max_pages):
        url = (
            "https://www.linkedin.com/voyager/api/graphql"
            f"?variables=(count:{page_size},start:{start})"
            "&queryId=voyagerJobsDashSavedJobPostingsByMember.f9ffd3a93b94f4e03e6a55301002cda7"
        )
        print(f"[LinkedIn Cookie] Page {page+1}, start={start}...")

        try:
            resp = http_requests.get(url, headers=headers, cookies=cookies, timeout=20)
        except Exception as e:
            print(f"[LinkedIn Cookie] Request error: {e}")
            break

        if resp.status_code in (401, 403):
            # Cookie expired — try re-login with credentials
            print("[LinkedIn Cookie] Session expired, attempting credential re-login...")
            new_li_at = _linkedin_login_for_cookie()
            if new_li_at:
                li_at = new_li_at
                cookies["li_at"] = li_at
                print("[LinkedIn Cookie] Re-login successful, retrying request...")
                try:
                    resp = http_requests.get(url, headers=headers, cookies=cookies, timeout=20)
                except Exception as e:
                    print(f"[LinkedIn Cookie] Retry error: {e}")
                    break
                if resp.status_code != 200:
                    return [], "LinkedIn session expired and re-login did not help"
            else:
                return [], "LinkedIn session expired — re-login failed. Verify from browser & retry."
        if resp.status_code != 200:
            # Try fallback REST endpoint
            url2 = (
                "https://www.linkedin.com/voyager/api/savesToDashJobPostings"
                f"?count={page_size}&q=savesToDashJobPostingsByMember&start={start}"
            )
            try:
                resp = http_requests.get(url2, headers=headers, cookies=cookies, timeout=20)
            except Exception as e:
                print(f"[LinkedIn Cookie] Fallback request error: {e}")
                break
            if resp.status_code in (401, 403):
                return [], "LinkedIn session expired — re-login failed. Verify from browser & retry."
            if resp.status_code != 200:
                print(f"[LinkedIn Cookie] API returned {resp.status_code}")
                break

        try:
            data = resp.json()
        except Exception:
            print(f"[LinkedIn Cookie] Invalid JSON response")
            break

        # Parse the response — LinkedIn uses "included" array with entity types
        included = data.get("included", data.get("elements", []))
        if not included and "data" in data:
            # GraphQL wrapper
            inner = data["data"]
            if isinstance(inner, dict):
                for v in inner.values():
                    if isinstance(v, dict) and "elements" in v:
                        included = v["elements"]
                        break
            if not included:
                included = data.get("included", [])

        # Build lookup maps for companies and job postings from included entities
        companies = {}
        postings = {}
        saved_meta = []  # saved-job wrapper entities with timestamps

        for item in included:
            urn = item.get("entityUrn") or item.get("$id") or ""
            t = item.get("$type", "")

            # Company / Organization
            if "Company" in t or "Organization" in t or "company" in urn:
                companies[urn] = item.get("name") or item.get("universalName") or "Unknown"

            # Job Posting
            if "JobPosting" in t or "jobPosting" in urn:
                postings[urn] = item

            # Saved-job wrapper (has savedAt timestamp)
            if "SavedJob" in t or "savedJob" in str(item.get("$recipeTypes", "")):
                saved_meta.append(item)

            # Also capture items that have 'title' + 'companyDetails' (direct posting)
            if item.get("title") and (item.get("companyDetails") or item.get("companyName")):
                postings[urn] = item

        # If we got saved_meta wrappers, extract job refs from them
        page_jobs = []
        if saved_meta:
            for sm in saved_meta:
                # Get savedAt timestamp for 30-day filter
                saved_at = sm.get("savedAt") or sm.get("createdAt") or 0
                if isinstance(saved_at, (int, float)) and saved_at > 1e12:
                    saved_at = saved_at / 1000  # ms to seconds
                if saved_at and saved_at > 0:
                    saved_dt = datetime.datetime.fromtimestamp(saved_at, tz=datetime.timezone.utc)
                    if saved_dt < cutoff:
                        continue  # Skip jobs saved more than max_days ago

                # Find the job posting reference
                jp_ref = sm.get("jobPosting") or sm.get("*jobPosting") or ""
                if isinstance(jp_ref, dict):
                    jp_ref = jp_ref.get("entityUrn") or jp_ref.get("$id") or ""
                jp = postings.get(jp_ref, {})
                if not jp and isinstance(jp_ref, str):
                    # Try partial match
                    for k, v in postings.items():
                        if jp_ref in k or k in jp_ref:
                            jp = v
                            break

                title = jp.get("title") or sm.get("title") or ""
                if not title:
                    continue

                # Company name
                company = "Unknown"
                comp_ref = jp.get("companyDetails", {}).get("*companyResolutionResult") or \
                           jp.get("companyDetails", {}).get("company") or \
                           jp.get("*company") or ""
                if isinstance(comp_ref, str) and comp_ref in companies:
                    company = companies[comp_ref]
                elif jp.get("companyName"):
                    company = jp["companyName"]
                else:
                    # Scan included for inline company names
                    for ci in included:
                        cn = ci.get("companyName")
                        if cn and ci.get("entityUrn") == jp.get("entityUrn"):
                            company = cn
                            break

                # Extract job ID from URN
                job_urn = jp.get("entityUrn") or jp_ref or ""
                m = re.search(r"(\d{8,})", job_urn)
                job_id = m.group(1) if m else ""
                job_url = f"https://www.linkedin.com/jobs/view/{job_id}/" if job_id else ""

                # JD text
                jd = ""
                desc = jp.get("description") or jp.get("descriptionText") or {}
                if isinstance(desc, dict):
                    jd = desc.get("text", "")
                elif isinstance(desc, str):
                    jd = desc

                page_jobs.append({
                    "role": title,
                    "company": company,
                    "url": job_url,
                    "linkedInId": f"li_{job_id}" if job_id else "",
                    "jd": jd[:4000],
                    "source": "LinkedIn",
                    "status": "saved",
                    "roleType": "Business Analyst",
                    "dateApplied": datetime.datetime.now().isoformat(),
                })
        else:
            # Fallback: parse postings directly (flat response format)
            for urn, jp in postings.items():
                title = jp.get("title", "")
                if not title:
                    continue

                # Check listedAt / repostedAt for 30-day filter
                listed = jp.get("listedAt") or jp.get("repostedAt") or 0
                if isinstance(listed, (int, float)) and listed > 1e12:
                    listed = listed / 1000
                if listed and listed > 0:
                    listed_dt = datetime.datetime.fromtimestamp(listed, tz=datetime.timezone.utc)
                    if listed_dt < cutoff:
                        continue

                company = "Unknown"
                comp_ref = jp.get("companyDetails", {}).get("*companyResolutionResult") or \
                           jp.get("*company") or ""
                if isinstance(comp_ref, str) and comp_ref in companies:
                    company = companies[comp_ref]
                elif jp.get("companyName"):
                    company = jp["companyName"]

                m = re.search(r"(\d{8,})", urn)
                job_id = m.group(1) if m else ""
                job_url = f"https://www.linkedin.com/jobs/view/{job_id}/" if job_id else ""

                jd = ""
                desc = jp.get("description") or jp.get("descriptionText") or {}
                if isinstance(desc, dict):
                    jd = desc.get("text", "")
                elif isinstance(desc, str):
                    jd = desc

                page_jobs.append({
                    "role": title,
                    "company": company,
                    "url": job_url,
                    "linkedInId": f"li_{job_id}" if job_id else "",
                    "jd": jd[:4000],
                    "source": "LinkedIn",
                    "status": "saved",
                    "roleType": "Business Analyst",
                    "dateApplied": datetime.datetime.now().isoformat(),
                })

        all_jobs.extend(page_jobs)
        print(f"[LinkedIn Cookie] Page {page+1}: {len(page_jobs)} jobs (total: {len(all_jobs)})")

        # Pagination: check if there are more
        paging = data.get("paging", {})
        if not paging and "data" in data:
            for v in data["data"].values():
                if isinstance(v, dict) and "paging" in v:
                    paging = v["paging"]
                    break
        total_results = paging.get("total", 0)
        raw_count = len(included)  # how many raw items LinkedIn returned (before our date filtering)
        # Stop only if: (a) LinkedIn returned NO raw items, or (b) we've fetched past the total count
        if raw_count == 0:
            print(f"[LinkedIn Cookie] No more raw items from API, stopping pagination")
            break
        if total_results > 0 and start + page_size >= total_results:
            print(f"[LinkedIn Cookie] Reached total ({total_results}), stopping pagination")
            break
        start += page_size
        _time.sleep(1)  # polite delay

    # ── Fetch JDs for jobs that don't have them yet ──
    if all_jobs:
        missing_jd = [j for j in all_jobs if not j.get("jd") and j.get("linkedInId")]
        if missing_jd:
            print(f"[LinkedIn Cookie] Fetching JDs for {len(missing_jd)} jobs...")
            for i, job in enumerate(missing_jd):
                job_id = job["linkedInId"].replace("li_", "")
                jd_url = (
                    f"https://www.linkedin.com/voyager/api/jobs/jobPostings/{job_id}"
                    f"?decorationId=com.linkedin.voyager.deco.jobs.web.shared.WebFullJobPosting-65"
                )
                try:
                    r = http_requests.get(jd_url, headers=headers, cookies=cookies, timeout=15)
                    if r.status_code == 200:
                        jdata = r.json()
                        desc = jdata.get("description") or jdata.get("descriptionText") or {}
                        if isinstance(desc, dict):
                            job["jd"] = desc.get("text", "")[:4000]
                        elif isinstance(desc, str):
                            job["jd"] = desc[:4000]
                except Exception as e:
                    print(f"[LinkedIn Cookie] JD fetch error for {job_id}: {e}")
                if i % 10 == 9:
                    _time.sleep(1)

    # Deduplicate by job URL
    seen = set()
    unique = []
    for j in all_jobs:
        key = j.get("url") or j.get("linkedInId") or j.get("role", "")
        if key in seen:
            continue
        seen.add(key)
        unique.append(j)

    jd_count = sum(1 for j in unique if j.get("jd"))
    print(f"[LinkedIn Cookie] Done — {len(unique)} unique saved jobs, {jd_count} with JD")
    return unique, None


# ─── LINKEDIN IMPORT ENDPOINT (receives jobs from bookmarklet) ──────────────

@app.route("/api/linkedin-import", methods=["POST"])
def linkedin_import_from_bookmarklet():
    """
    Receive saved jobs POSTed from the bookmarklet running in the user's browser.
    The bookmarklet scrapes LinkedIn saved-jobs page + fetches JDs via Voyager API,
    then sends the results here for server-side storage.
    """
    import time as _ts
    import traceback

    try:
        data = request.json or {}
        secret = data.get("secret", "")
        if secret != AGENT_CRON_SECRET:
            return jsonify({"error": "Unauthorized"}), 401

        jobs = data.get("jobs", [])
        if not jobs:
            return jsonify({"status": "ok", "jobs_added": 0, "jobs_skipped": 0, "message": "No jobs received"})

        print(f"[LinkedIn Import] Received {len(jobs)} jobs from bookmarklet")

        sb = get_supabase()
        if not sb:
            return jsonify({"status": "error", "error": "Supabase connection failed"}), 500

        ex_res = sb.table("jobs").select("url,linkedInId").execute()
        ex_urls = {(j.get("url") or "").split("?")[0].rstrip("/") for j in (ex_res.data or []) if j.get("url")}
        ex_li_ids = {j.get("linkedInId") for j in (ex_res.data or []) if j.get("linkedInId")}

        to_insert = []
        skipped = 0
        for lj in jobs:
            curl = (lj.get("url") or "").split("?")[0].rstrip("/")
            li_id = lj.get("linkedInId", "")
            if (curl and curl in ex_urls) or (li_id and li_id in ex_li_ids):
                skipped += 1
                continue
            if curl: ex_urls.add(curl)
            if li_id: ex_li_ids.add(li_id)
            to_insert.append({
                "id": li_id or f"li_{int(_ts.time()*1000)}_{len(to_insert)}",
                "role": lj.get("role", ""),
                "company": lj.get("company", "Unknown"),
                "url": curl,
                "linkedInId": li_id,
                "jd": (lj.get("jd") or "")[:8000],
                "status": "saved",
                "source": "LinkedIn",
                "roleType": lj.get("roleType", ""),
                "dateApplied": lj.get("dateApplied", datetime.datetime.now().isoformat()),
                "datePosted": lj.get("datePosted", ""),
                "companyLogo": lj.get("companyLogo", ""),
            })

        if to_insert:
            BATCH = 50
            for i in range(0, len(to_insert), BATCH):
                sb.table("jobs").upsert(to_insert[i:i+BATCH], on_conflict="id").execute()

        added = len(to_insert)
        print(f"[LinkedIn Import] Done — {added} new, {skipped} duplicates")

        return jsonify({
            "status": "ok",
            "jobs_added": added,
            "jobs_skipped": skipped,
            "message": f"Imported {added} jobs ({skipped} duplicates skipped).",
        })

    except Exception as e:
        print(f"[LinkedIn Import] Error: {e}")
        traceback.print_exc()
        return jsonify({"status": "error", "error": str(e)}), 500


# ─── LINKEDIN SAVED JOBS ONLY ENDPOINT ──────────────────────────────────────

@app.route("/api/linkedin-saved-jobs", methods=["POST"])
def linkedin_saved_jobs_only():
    """
    Scrape ONLY LinkedIn saved jobs (last 30 days) via credentials / li_at cookie.
    Runs synchronously — returns the result directly (no background thread).
    """
    import time as _ts
    import traceback

    try:
        # Step 1: Get li_at cookie (stored or via credential login)
        li_at, li_src = _get_or_login_li_at()
        if not li_at:
            email = _get_linkedin_email()
            if email:
                return jsonify({
                    "status": "error",
                    "error": "LinkedIn login failed — may need security verification. Log in from browser, then retry.",
                    "jobs_added": 0, "jobs_skipped": 0,
                })
            return jsonify({
                "status": "error",
                "error": "No LinkedIn credentials set. Go to Settings and enter your LinkedIn email + password.",
                "jobs_added": 0, "jobs_skipped": 0,
            })

        # Step 2: Scrape saved jobs
        print(f"[LinkedIn-Only] Scraping saved jobs (auth via {li_src})...")
        try:
            li_jobs, li_err = linkedin_scrape_saved_jobs_via_cookie(max_days=30)
        except Exception as e:
            print(f"[LinkedIn-Only] Scrape error: {e}")
            traceback.print_exc()
            return jsonify({"status": "error", "error": f"Scrape error: {str(e)}", "jobs_added": 0, "jobs_skipped": 0})

        if li_err:
            return jsonify({"status": "error", "error": li_err, "jobs_added": 0, "jobs_skipped": 0})

        if not li_jobs:
            return jsonify({"status": "ok", "message": "No saved jobs found in the last 30 days.", "jobs_added": 0, "jobs_skipped": 0})

        # Step 3: Sync to Supabase (dedup by URL + linkedInId)
        sb = get_supabase()
        if not sb:
            return jsonify({"status": "error", "error": "Supabase connection failed", "jobs_added": 0, "jobs_skipped": 0})

        ex_res = sb.table("jobs").select("url,linkedInId").execute()
        ex_urls = {(j.get("url") or "").split("?")[0] for j in (ex_res.data or []) if j.get("url")}
        ex_li_ids = {j.get("linkedInId") for j in (ex_res.data or []) if j.get("linkedInId")}

        to_insert = []
        skipped = 0
        for lj in li_jobs:
            curl = (lj.get("url") or "").split("?")[0]
            li_id = lj.get("linkedInId", "")
            if (curl and curl in ex_urls) or (li_id and li_id in ex_li_ids):
                skipped += 1
                continue
            if curl: ex_urls.add(curl)
            if li_id: ex_li_ids.add(li_id)
            to_insert.append({
                "id": li_id or f"li_{int(_ts.time()*1000)}_{len(to_insert)}",
                "role": lj.get("role", ""),
                "company": lj.get("company", ""),
                "url": curl,
                "linkedInId": li_id,
                "jd": (lj.get("jd") or "")[:8000],
                "status": "saved",
                "source": "LinkedIn",
                "roleType": "Business Analyst",
                "dateApplied": lj.get("dateApplied", datetime.datetime.now().isoformat()),
            })

        if to_insert:
            BATCH = 50
            for i in range(0, len(to_insert), BATCH):
                sb.table("jobs").upsert(to_insert[i:i+BATCH], on_conflict="id").execute()

        added = len(to_insert)
        print(f"[LinkedIn-Only] Done — {added} new, {skipped} duplicates")

        return jsonify({
            "status": "ok",
            "message": f"Imported {added} saved jobs from LinkedIn ({skipped} duplicates skipped).",
            "jobs_added": added,
            "jobs_skipped": skipped,
            "auth_method": li_src,
        })

    except Exception as e:
        print(f"[LinkedIn-Only] Unexpected error: {e}")
        traceback.print_exc()
        return jsonify({"status": "error", "error": f"Server error: {str(e)}", "jobs_added": 0, "jobs_skipped": 0}), 500


@app.route("/api/agent/full-run", methods=["POST"])
def agent_full_run():
    """
    Button-triggered full pipeline (no Selenium — uses lightweight HTTP scrapers):
      1. Discover jobs from MCF, Workable, LinkedIn Guest/Public, MCF Extended
      2. Deduplicate & sync new jobs to Supabase
      3. AI score all unscored jobs
      4. Generate resume + cover letter for scored >= 5
      5. Email / WhatsApp notification
    """
    def bg():
        with app.app_context():
            summary = {"scraped": 0, "skipped": 0, "scored": 0, "docs": 0, "error": None}
            li_saved_count = 0

            # ── Step 0: Scrape LinkedIn saved jobs (credentials → cookie → Voyager API) ──
            li_at, li_src = _get_or_login_li_at()
            if li_at:
                print(f"[FullRun] Scraping LinkedIn saved jobs (auth via {li_src})...")
                try:
                    li_jobs, li_err = linkedin_scrape_saved_jobs_via_cookie(max_days=30)
                    if li_err:
                        print(f"[FullRun] LinkedIn cookie error: {li_err}")
                        summary["error"] = li_err
                    elif li_jobs:
                        # Sync LinkedIn saved jobs to Supabase
                        sb_li = get_supabase()
                        if sb_li:
                            ex_res = sb_li.table("jobs").select("url,linkedInId").execute()
                            ex_urls = {(j.get("url") or "").split("?")[0] for j in (ex_res.data or []) if j.get("url")}
                            ex_li_ids = {j.get("linkedInId") for j in (ex_res.data or []) if j.get("linkedInId")}
                            import time as _time_li
                            to_insert = []
                            for lj in li_jobs:
                                curl = (lj.get("url") or "").split("?")[0]
                                li_id = lj.get("linkedInId", "")
                                if (curl and curl in ex_urls) or (li_id and li_id in ex_li_ids):
                                    summary["skipped"] += 1
                                    continue
                                if curl: ex_urls.add(curl)
                                if li_id: ex_li_ids.add(li_id)
                                to_insert.append({
                                    "id": li_id or f"li_{int(_time_li.time()*1000)}_{len(to_insert)}",
                                    "role": lj.get("role", ""),
                                    "company": lj.get("company", ""),
                                    "url": curl,
                                    "linkedInId": li_id,
                                    "jd": (lj.get("jd") or "")[:8000],
                                    "status": "saved",
                                    "source": "LinkedIn",
                                    "roleType": "Business Analyst",
                                    "dateApplied": lj.get("dateApplied", datetime.datetime.now().isoformat()),
                                })
                            if to_insert:
                                BATCH = 50
                                for i in range(0, len(to_insert), BATCH):
                                    sb_li.table("jobs").upsert(to_insert[i:i+BATCH], on_conflict="id").execute()
                            li_saved_count = len(to_insert)
                            summary["scraped"] += li_saved_count
                            print(f"[FullRun] LinkedIn saved: {li_saved_count} new, {summary['skipped']} duplicates")
                except Exception as e:
                    print(f"[FullRun] LinkedIn cookie scrape error: {e}")
            else:
                email = _get_linkedin_email()
                if email:
                    print("[FullRun] LinkedIn credential login failed — skipping saved jobs")
                    summary["error"] = "LinkedIn login failed. Log in via browser & retry, or paste li_at cookie."
                else:
                    print("[FullRun] No LinkedIn credentials or cookie set — skipping saved jobs")

            # ── Step 1: Discover jobs from all 5 HTTP scrapers ──
            print("[FullRun] Discovering jobs from 5 platforms (no login needed)...")
            try:
                import time as _time_disc
                P = get_active_profile()
                kw_parts = P.get("headline", "Business Analyst").split("|")
                kw = kw_parts[0].strip() if kw_parts else "Business Analyst"

                all_scrapers = {
                    "mycareersfuture": _scrape_mycareersfuture,
                    "workable": _scrape_workable,
                    "linkedin_guest": _scrape_linkedin_guest,
                    "linkedin_public": _scrape_linkedin_public,
                    "mcf_extended": _scrape_mcf_extended,
                }
                disc_jobs = []
                scraper_details = {}
                with ThreadPoolExecutor(max_workers=5) as tex:
                    futs = {tex.submit(fn, kw, "Singapore", 14): name for name, fn in all_scrapers.items()}
                    for fut in as_completed(futs):
                        name = futs[fut]
                        try:
                            batch = fut.result(timeout=30)
                            # Some scrapers return (jobs, error), some return just jobs
                            if isinstance(batch, tuple):
                                batch = batch[0] or []
                            disc_jobs.extend(batch)
                            scraper_details[name] = len(batch)
                            print(f"[FullRun] {name}: {len(batch)} jobs")
                        except Exception as e:
                            scraper_details[name] = 0
                            print(f"[FullRun] {name}: error — {str(e)[:60]}")

                print(f"[FullRun] Total discovered: {len(disc_jobs)} jobs from {len(scraper_details)} platforms")

                # Deduplicate by URL + title|company
                seen = set()
                unique_disc = []
                for dj in disc_jobs:
                    curl = (dj.get("url") or "").split("?")[0]
                    tc = f"{dj.get('role','').lower().strip()}|{dj.get('company','').lower().strip()}"
                    if (curl and curl in seen) or tc in seen:
                        continue
                    if curl: seen.add(curl)
                    seen.add(tc)
                    unique_disc.append(dj)

                # Sync to Supabase, skipping existing
                sb_disc = get_supabase()
                if sb_disc and unique_disc:
                    ex_res = sb_disc.table("jobs").select("url").execute()
                    ex_urls = {(j.get("url") or "").split("?")[0] for j in (ex_res.data or []) if j.get("url")}
                    to_insert = []
                    for dj in unique_disc:
                        curl = (dj.get("url") or "").split("?")[0]
                        if not curl or curl in ex_urls:
                            summary["skipped"] += 1
                            continue
                        ex_urls.add(curl)
                        to_insert.append({
                            "id": f"disc_{int(_time_disc.time()*1000)}_{len(to_insert)}",
                            "role": dj.get("role", ""),
                            "company": dj.get("company", ""),
                            "url": curl,
                            "jd": (dj.get("jd") or "")[:8000],
                            "status": "wishlist",
                            "source": dj.get("platform", dj.get("source", "Discovery")),
                            "roleType": "Business Analyst",
                            "dateApplied": datetime.datetime.now().isoformat(),
                        })
                    # Batch upsert
                    if to_insert:
                        BATCH = 50
                        for i in range(0, len(to_insert), BATCH):
                            sb_disc.table("jobs").upsert(to_insert[i:i+BATCH], on_conflict="id").execute()
                    summary["scraped"] = len(to_insert)
                    print(f"[FullRun] Synced: {len(to_insert)} new, {summary['skipped']} duplicates")

                # Build platform breakdown for email
                li_line = f"<li><strong>LinkedIn Saved Jobs: {li_saved_count} new</strong></li>" if li_saved_count else ""
                platform_lines = li_line + "".join(
                    f"<li>{n}: {c} jobs</li>" for n, c in scraper_details.items()
                )
                send_email(
                    f"🔍 Discovered {summary['scraped']} new jobs from 5 platforms",
                    f"<h2>Job Discovery Complete</h2>"
                    f"<p><strong>{summary['scraped']}</strong> new jobs added, "
                    f"<strong>{summary['skipped']}</strong> duplicates skipped.</p>"
                    f"<ul>{platform_lines}</ul>"
                    f"<p><a href='https://job-hunt-app-r7my.onrender.com'>Open tracker →</a></p>"
                )
            except Exception as e:
                summary["error"] = str(e)
                print(f"[FullRun] Discovery error: {e}")
                send_email("⚠️ Job Discovery Failed", f"<h2>Error</h2><p>{e}</p>")

            # ── Step 2 + 3: Score + generate docs ──
            sb = get_supabase()
            if not sb:
                return
            try:
                res  = sb.table("jobs").select("*").execute()
                jobs = [j for j in (res.data or []) if not j.get("isDemo")]
                to_run = [
                    j for j in jobs
                    if (j.get("jd") and j.get("aiScore") is None) or
                       (j.get("aiScore", 0) >= 5 and not j.get("resume_docx_b64"))
                ]
                if to_run:
                    print(f"[FullRun] Scoring/doc-gen for {len(to_run)} jobs...")
                    agent_run(to_run, trigger="manual")
                else:
                    send_email(
                        f"✅ Extract Latest Jobs — {summary['scraped']} new jobs added",
                        f"""<html><body style="font-family:sans-serif;padding:20px;">
                        <h2>🔍 Job Discovery Complete</h2>
                        <p><strong>{summary['scraped']}</strong> new jobs added to tracker</p>
                        <p><strong>{summary['skipped']}</strong> duplicates skipped</p>
                        <p>All jobs already scored — no new processing needed.</p>
                        <p><a href="https://job-hunt-app-r7my.onrender.com">Open your tracker →</a></p>
                        </body></html>"""
                    )
            except Exception as e:
                print(f"[FullRun] Agent error: {e}")

    threading.Thread(target=bg, daemon=True).start()
    return jsonify({"status": "started"})

@app.route("/api/agent/cron", methods=["POST", "GET"])
def agent_cron():
    """Daily cron trigger — discovers jobs via HTTP scrapers then runs AI agent pipeline."""
    secret = request.args.get("secret") or (request.json or {}).get("secret", "")
    if secret != AGENT_CRON_SECRET:
        return jsonify({"error": "Unauthorized"}), 401

    def bg():
        with app.app_context():
            # Step 1: Discover jobs from all platforms (lightweight HTTP — no Chrome)
            print("[Cron] Discovering jobs from all platforms...")
            try:
                import time as _time_cron
                P = get_active_profile()
                kw_parts = P.get("headline", "Business Analyst").split("|")
                kw = kw_parts[0].strip() if kw_parts else "Business Analyst"
                cron_scrapers = {
                    "mycareersfuture": _scrape_mycareersfuture,
                    "workable": _scrape_workable,
                    "linkedin_guest": _scrape_linkedin_guest,
                    "linkedin_public": _scrape_linkedin_public,
                    "mcf_extended": _scrape_mcf_extended,
                }
                disc_jobs = []
                with ThreadPoolExecutor(max_workers=5) as tex:
                    futs = {tex.submit(fn, kw, "Singapore", 14): n for n, fn in cron_scrapers.items()}
                    for fut in as_completed(futs):
                        try:
                            batch = fut.result(timeout=30)
                            if isinstance(batch, tuple):
                                batch = batch[0] or []
                            disc_jobs.extend(batch)
                        except Exception:
                            pass

                sb = get_supabase()
                added = 0
                if sb and disc_jobs:
                    ex_urls = {(j.get("url") or "").split("?")[0]
                               for j in (sb.table("jobs").select("url").execute().data or [])
                               if j.get("url")}
                    to_ins = []
                    for dj in disc_jobs:
                        curl = (dj.get("url") or "").split("?")[0]
                        if not curl or curl in ex_urls:
                            continue
                        ex_urls.add(curl)
                        to_ins.append({
                            "id": f"disc_{int(_time_cron.time()*1000)}_{len(to_ins)}",
                            "role": dj.get("role", ""),
                            "company": dj.get("company", ""),
                            "url": curl,
                            "jd": (dj.get("jd") or "")[:8000],
                            "status": "wishlist",
                            "source": dj.get("platform", dj.get("source", "Discovery")),
                            "roleType": "Business Analyst",
                            "dateApplied": datetime.datetime.now().isoformat(),
                        })
                    if to_ins:
                        BATCH = 50
                        for i in range(0, len(to_ins), BATCH):
                            sb.table("jobs").upsert(to_ins[i:i+BATCH], on_conflict="id").execute()
                    added = len(to_ins)
                print(f"[Cron] Discovery: {added} new jobs from {len(disc_jobs)} discovered")
            except Exception as e:
                print(f"[Cron] Discovery error: {e}")

            # Step 2: Run AI agent on all pending jobs
            sb = get_supabase()
            if not sb:
                return
            try:
                res  = sb.table("jobs").select("*").execute()
                jobs = [j for j in (res.data or []) if not j.get("isDemo")]
                to_run = [
                    j for j in jobs
                    if (j.get("jd") and j.get("aiScore") is None) or
                       (j.get("aiScore", 0) >= 5 and not j.get("resume_docx_b64"))
                ]
                if to_run:
                    agent_run(to_run, trigger="cron")
                else:
                    print("[Cron] No jobs to process")
                    send_email("✅ Daily Cron — Nothing to Process",
                               "<h2>Daily Job Agent</h2><p>All jobs already scored and docs generated.</p>")
            except Exception as e:
                print(f"[Cron] Agent error: {e}")

    threading.Thread(target=bg, daemon=True).start()
    return jsonify({"status": "started", "message": "Job discovery + agent pipeline running"})


@app.route("/api/test-notifications", methods=["POST"])
def test_notifications():
    ok = send_whatsapp("🤖 Job Agent test — WhatsApp connected ✅")
    return jsonify({
        "whatsapp": "✅ sent" if ok else "❌ not configured — check Twilio credentials and NOTIFICATION_PHONE"
    })


if __name__ == "__main__":
    # Ensure settings table exists at startup
    try:
        ensure_settings_table()
    except Exception as e:
        print(f"[Startup] Settings table check: {e}")
    port = int(os.environ.get("PORT", 5000))
    app.run(host="0.0.0.0", port=port, debug=False)
